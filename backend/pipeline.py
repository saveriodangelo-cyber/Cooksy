from __future__ import annotations

import os
import re
import json
import time
import random
import string
import math
import urllib.request
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
import uuid
import logging
import sys

# Consenti l'esecuzione diretta del file inserendo il root progetto nel path
_ROOT = Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from backend.utils import project_root

try:
    from backend.app_logging import get_logger, log_event, truncate_text
except Exception:
    def get_logger(*args: Any, **kwargs: Any) -> logging.Logger:
        return logging.getLogger()
    def log_event(*args: Any, **kwargs: Any) -> None:
        pass
    def truncate_text(text: Any, max_len: int = 500) -> str:
        return str(text)[:max_len]

try:
    from backend.prices_db import PricesDB
except Exception:
    PricesDB = None

try:
    from backend.nutrition_db import NutritionDB
except Exception:
    NutritionDB = None

try:
    from backend.archive_db import ArchiveDB
except Exception:
    ArchiveDB = None

try:
    from backend import price_engine  # type: ignore
except Exception:
    price_engine = None

try:
    from backend import nutrition_engine
except Exception:
    nutrition_engine = None

try:
    from backend import allergens
except Exception:
    allergens = None

try:
    from backend import equipment  # type: ignore
except Exception:
    equipment = None

try:
    from backend import ai_cloud
except Exception:
    ai_cloud = None

try:
    from backend import cloud_ai as cloud_ai_settings
except Exception:
    cloud_ai_settings = None

try:
    from backend.pdf_export import (
        export_recipe_pdf,
        list_templates as _get_pdf_templates,
        render_html_template as _render_html_template,
    )
except Exception:
    export_recipe_pdf = None
    _get_pdf_templates = None
    _render_html_template = None

try:
    from backend.docx_export import export_recipe_docx
except Exception:
    export_recipe_docx = None

try:
    from backend.parser_engine import parse_recipe_text
except Exception:
    def parse_recipe_text(text: str) -> Tuple[Dict[str, Any], List[str]]:
        return {"title": "Ricetta", "ingredients": [], "steps": []}, []

try:
    from backend import ocr_engines
except Exception:
    ocr_engines = None


# ------------------------------------------------------------
# CONFIG
# ------------------------------------------------------------
OLLAMA_URL = os.environ.get("RIGENERA_OLLAMA_URL") or os.environ.get("RICETTEPDF_OLLAMA_URL") or "http://127.0.0.1:11434"
OLLAMA_MODEL = (os.environ.get("RICETTEPDF_OLLAMA_MODEL") or os.environ.get("RIGENERA_OLLAMA_MODEL") or "llama3.2").strip()

# Validate Ollama URL for security: only allow localhost/127.0.0.1 (local) or HTTPS (secure remote)
_OLLAMA_LOWER = OLLAMA_URL.lower().strip()
if not (_OLLAMA_LOWER.startswith("http://127.0.0.1") or _OLLAMA_LOWER.startswith("http://localhost") or _OLLAMA_LOWER.startswith("https://")):
    raise ValueError(
        f"Ollama URL must be localhost (http://localhost:11434) or HTTPS (https://...), got: {OLLAMA_URL}. "
        "Set RICETTEPDF_OLLAMA_URL environment variable correctly for security."
    )
MAX_SOURCE_CHARS = 20000
DEFAULT_EXPORT_PDF_SIZE = "A4"

_IMG_EXT = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"}
_PDF_EXT = {".pdf"}
_DOCX_EXT = {".docx"}
_TXT_EXT = {".txt", ".log", ".md"}


# ------------------------------------------------------------
# Progress helper
# ------------------------------------------------------------
def _p_set(progress_cb: Any, pct: int, stage: str, msg: str) -> None:
    if progress_cb is None:
        return
    try:
        if callable(progress_cb):
            progress_cb(int(pct), str(msg), str(stage))
            return
        if hasattr(progress_cb, "set"):
            progress_cb.set(int(pct), str(stage), str(msg))
            return
        if isinstance(progress_cb, dict):
            progress_cb["pct"] = int(pct)
            progress_cb["stage"] = str(stage)
            progress_cb["msg"] = str(msg)
    except Exception:
        pass


class _ProgressPlan:
    def __init__(self, progress_cb: Any, steps: List[Tuple[str, float]]) -> None:
        self._progress = progress_cb
        total = sum(w for _, w in steps) if steps else 0.0
        total = total if total > 0 else 1.0
        self._ranges: Dict[str, Tuple[float, float]] = {}
        current = 0.0
        for name, weight in steps:
            span = (float(weight) / total) * 100.0
            self._ranges[name] = (current, current + span)
            current += span
        self._last_pct = 0.0

    def set(self, name: str, pct: float, msg: str) -> None:
        if self._progress is None:
            return
        start, end = self._ranges.get(name, (self._last_pct, 100.0))
        pct_clamped = max(0.0, min(100.0, float(pct)))
        mapped = start + ((end - start) * (pct_clamped / 100.0))
        if mapped < self._last_pct:
            mapped = self._last_pct
        self._last_pct = mapped
        _p_set(self._progress, int(round(mapped)), name, msg)

    def done(self, name: str, msg: str) -> None:
        self.set(name, 100.0, msg)


# ------------------------------------------------------------
# Text extraction
# ------------------------------------------------------------
def _safe_read_text(path: str) -> str:
    try:
        return Path(path).read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


def _extract_text_from_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        parts: List[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join([p for p in parts if p]).strip()
    except Exception:
        return ""


def _extract_text_from_docx(path: str) -> str:
    try:
        from docx import Document
    except Exception:
        return ""

    def _read_docx() -> str:
        doc = Document(path)
        parts: List[str] = []
        char_count = 0
        max_chars = 60000
        max_rows = 2000
        row_count = 0

        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
                char_count += len(t)
                if char_count >= max_chars:
                    break

        if char_count < max_chars:
            for table in doc.tables:
                for row in table.rows:
                    row_count += 1
                    if row_count > max_rows or char_count >= max_chars:
                        break
                    cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                    if cells:
                        line = " | ".join(cells)
                        parts.append(line)
                        char_count += len(line)
                if row_count > max_rows or char_count >= max_chars:
                    break

        return "\n".join(parts).strip()

    try:
        from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
    except Exception:
        ThreadPoolExecutor = None  # type: ignore[assignment]
        FuturesTimeoutError = None  # type: ignore[assignment]

    if ThreadPoolExecutor is None:
        try:
            return _read_docx()
        except Exception:
            return ""

    try:
        with ThreadPoolExecutor(max_workers=1) as ex:
            fut = ex.submit(_read_docx)
            return fut.result(timeout=45)
    except Exception as e:
        if FuturesTimeoutError is not None and isinstance(e, FuturesTimeoutError):
            return ""
        return ""


def _ocr_with_tesseract(image_paths: List[str], lang: str) -> str:
    try:
        import pytesseract
        from PIL import Image
    except Exception:
        return ""
    parts: List[str] = []
    for p in image_paths:
        try:
            parts.append(pytesseract.image_to_string(Image.open(p), lang=lang))
        except Exception:
            continue
    return "\n".join(parts).strip()


def _extract_text_from_images(
    image_paths: List[str],
    lang: str = "ita",
    ocr_strategy: str = "multi",
    progress_cb: Any = None,
) -> Tuple[str, Dict[str, Any]]:
    report: Dict[str, Any] = {"engine": None, "text_len": 0}
    if not image_paths:
        return "", report

    strat = (ocr_strategy or "").strip().lower()
    if strat in {"multi", "auto"} and ocr_engines is not None:
        try:
            text, rep = ocr_engines.ocr_images_combined(image_paths, lang=lang, progress_cb=progress_cb)
            report["engine"] = rep.get("selected") if isinstance(rep, dict) else "multi"
            report["text_len"] = len(text or "")
            report["report"] = rep
            return text, report
        except Exception:
            pass

    text = _ocr_with_tesseract(image_paths, lang=lang)
    report["engine"] = "tesseract" if text else None
    report["text_len"] = len(text or "")
    return text, report


def extract_text(path: str) -> str:
    p = Path(path)
    suffix = p.suffix.lower()

    if suffix in _TXT_EXT:
        return _safe_read_text(str(p))
    if suffix in _PDF_EXT:
        return _extract_text_from_pdf(str(p))
    if suffix in _DOCX_EXT:
        return _extract_text_from_docx(str(p))
    if suffix in _IMG_EXT:
        text, _rep = _extract_text_from_images([str(p)])
        return text
    return ""


def extract_text_from_paths(
    paths: List[str],
    options: Optional[Dict[str, Any]] = None,
    progress_cb: Any = None,
) -> Tuple[str, Dict[str, Any]]:
    opts = options or {}
    lang = str(opts.get("ocr_lang") or "ita")
    ocr_strategy = str(opts.get("ocr_strategy") or "multi")

    texts: List[str] = []
    image_paths: List[str] = []

    for p in paths:
        ext = Path(p).suffix.lower()
        if ext in _IMG_EXT:
            image_paths.append(p)
            continue
        t = extract_text(p)
        if t:
            texts.append(f"\n\n# FILE: {Path(p).name}\n{t}")

    ocr_report: Dict[str, Any] = {}
    if image_paths:
        timg, rep = _extract_text_from_images(image_paths, lang=lang, ocr_strategy=ocr_strategy, progress_cb=progress_cb)
        if timg:
            texts.append(f"\n\n# IMAGES OCR\n{timg}")
        ocr_report = rep

    combined = "\n".join(texts).strip()
    truncated = False
    if len(combined) > MAX_SOURCE_CHARS:
        combined = combined[:MAX_SOURCE_CHARS] + "\n...[TRONCATO]"
        truncated = True
    debug = {
        "ocr_engine_used": ocr_report.get("engine"),
        "ocr_text_len": ocr_report.get("text_len", 0),
        "text_truncated": truncated,
        "text_len": len(combined),
    }
    return combined, debug


# ------------------------------------------------------------
# Local AI (Ollama)
# ------------------------------------------------------------
def _ask_ollama_json(prompt: str, model: str = OLLAMA_MODEL) -> str:
    payload = {"model": model, "prompt": prompt, "stream": False}
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        f"{OLLAMA_URL}/api/generate",
        data=data,
        headers={"Content-Type": "application/json"},
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            return json.loads(resp.read().decode("utf-8")).get("response", "")
    except Exception as e:
        print(f"Ollama Error: {e}")
        return ""


def standard_recipe_extraction(text: str) -> str:
    prompt = (
        "Sei un assistente chef. Estrai la ricetta dal testo seguente.\n"
        "Rispondi SOLO con un JSON valido.\n"
        "Schema JSON richiesto:\n"
        "{\n"
        '  "title": "...",\n'
        '  "category": "Primi/Secondi/Dolci/...",\n'
        '  "servings": 4,\n'
        '  "difficulty": "bassa|media|alta",\n'
        '  "prep_time_min": 0,\n'
        '  "cook_time_min": 0,\n'
        '  "total_time_min": 0,\n'
        '  "ingredients": [ {\n'
        '    "name": "...", "qty": 0, "unit": "g|kg|ml|l|pz|cucchiai"\n'
        "  } ],\n"
        '  "steps": [ "Passo 1...", "Passo 2..." ]\n'
        "}\n\n"
        f"TESTO ORIGINALE:\n{text[:MAX_SOURCE_CHARS]}"
    )
    return _ask_ollama_json(prompt)


def _is_empty_value(val: Any) -> bool:
    if val is None:
        return True
    if isinstance(val, str):
        s = val.strip()
        if not s:
            return True
        return s.lower() in {"none", "null", "n/d", "n.d", "n.d.", "nd", "n.a", "n/a"}
    if isinstance(val, (list, tuple, set, dict)):
        return len(val) == 0
    return False


def _is_zero_like(val: Any) -> bool:
    if val is None:
        return False
    if isinstance(val, (int, float)):
        return float(val) == 0.0
    if isinstance(val, str):
        return val.strip().replace(",", ".") in {"0", "0.0", "0.00"}
    return False


def _collect_missing_fields(recipe: Dict[str, Any], missing: Optional[List[str]] = None) -> List[str]:
    fields = set()
    for item in (missing or []):
        if not isinstance(item, str):
            continue
        if ":" in item:
            continue
        fields.add(item)

    if _is_empty_value(recipe.get("title")):
        fields.add("title")
    if _is_empty_value(recipe.get("category")):
        fields.add("category")
    if _is_empty_value(recipe.get("servings")):
        fields.add("servings")
    if _is_empty_value(recipe.get("difficulty")):
        fields.add("difficulty")
    if _is_empty_value(recipe.get("ingredients")):
        fields.add("ingredients")
    if _is_empty_value(recipe.get("steps")):
        fields.add("steps")

    if (
        _is_empty_value(recipe.get("prep_time_min"))
        and _is_empty_value(recipe.get("cook_time_min"))
        and _is_empty_value(recipe.get("total_time_min"))
    ):
        fields.add("time")

    diet_flags = recipe.get("diet_flags")
    has_diet_flags = isinstance(diet_flags, dict) and any(bool(v) for v in diet_flags.values())
    if _is_empty_value(recipe.get("diet_text")) and not has_diet_flags:
        fields.add("diets")

    if _is_empty_value(recipe.get("conservazione")):
        fields.add("conservazione")
    if _is_empty_value(recipe.get("allergens_text")) and _is_empty_value(recipe.get("allergens")):
        fields.add("allergens")
    if _is_empty_value(recipe.get("vino_descrizione")) and _is_empty_value(recipe.get("wine_pairing")):
        fields.add("wine_pairing")
    if _is_empty_value(recipe.get("vino_temperatura_servizio")):
        fields.add("vino_temperatura_servizio")
    if _is_empty_value(recipe.get("vino_regione")):
        fields.add("vino_regione")
    if _is_empty_value(recipe.get("vino_annata")):
        fields.add("vino_annata")
    if _is_empty_value(recipe.get("vino_motivo_annata")):
        fields.add("vino_motivo_annata")
    if _is_empty_value(recipe.get("equipment_text")) and _is_empty_value(recipe.get("attrezzature_generiche")):
        fields.add("equipment")
    if _is_empty_value(recipe.get("presentazione_impiattamento")):
        fields.add("presentazione")
    if _is_empty_value(recipe.get("stagionalita")):
        fields.add("stagionalita")

    cost_lines = recipe.get("cost_lines")
    has_cost_lines = isinstance(cost_lines, list) and any(isinstance(x, dict) for x in cost_lines)
    has_cost_values = False
    scarto_missing = False
    if isinstance(cost_lines, list):
        for row in cost_lines:
            if not isinstance(row, dict):
                continue
            scarto_val = row.get("scarto") or row.get("scarto_pct") or row.get("waste_pct")
            if _is_empty_value(scarto_val):
                scarto_missing = True
            for key in (
                "prezzo_calcolato",
                "prezzo_kg_ud",
                "prezzo_alimento_acquisto",
                "prezzo_acquisto",
                "price_value",
                "cost",
            ):
                if not _is_empty_value(row.get(key)) and not _is_zero_like(row.get(key)):
                    has_cost_values = True
                    break
            if has_cost_values:
                break
    cost_missing = (not has_cost_lines) or (not has_cost_values)
    if scarto_missing:
        cost_missing = True
    if _is_empty_value(recipe.get("spesa_totale_ricetta")) or _is_zero_like(recipe.get("spesa_totale_ricetta")):
        cost_missing = True
    if _is_empty_value(recipe.get("spesa_totale_acquisto")) or _is_zero_like(recipe.get("spesa_totale_acquisto")):
        cost_missing = True
    if _is_empty_value(recipe.get("spesa_per_porzione")) or _is_zero_like(recipe.get("spesa_per_porzione")):
        cost_missing = True
    if cost_missing:
        fields.add("costs")

    nt = recipe.get("nutrition_table")
    required_nut = [
        "energia",
        "carboidrati_totali",
        "di_cui_zuccheri",
        "grassi_totali",
        "di_cui_saturi",
        "monoinsaturi",
        "polinsaturi",
        "proteine_totali",
        "colesterolo",
        "fibre",
        "sodio",
    ]

    def _block_missing(block: Any) -> bool:
        if not isinstance(block, dict):
            return True
        for key in required_nut:
            v = block.get(key)
            if _is_empty_value(v) or _is_zero_like(v):
                return True
        return False

    if not isinstance(nt, dict) or _block_missing(nt.get("100g")) or _block_missing(nt.get("totale")):
        fields.add("nutrition")

    return sorted(fields)


_TEMPLATE_REQUIRED_TEXT_FIELDS = [
    "titolo",
    "porzioni",
    "tempo_totale",
    "tempo_dettaglio",
    "tempo_preparazione",
    "tempo_cottura",
    "difficolta",
    "ingredienti_blocco",
    "procedimento_blocco",
    "conservazione",
    "diete_scelta_alimentare",
    "diete_cliniche",
    "diete_culturali",
    "diete_stile",
    "allergeni_elenco",
    "vino_descrizione",
    "vino_temperatura_servizio",
    "vino_regione",
    "vino_annata",
    "vino_motivo_annata",
    "attrezzature_specifiche",
    "attrezzature_generiche",
    "attrezzature_pasticceria",
    "presentazione_impiattamento",
    "stagionalita",
]

_TEMPLATE_REQUIRED_NUM_FIELDS = [
    "energia_100g",
    "energia_totale",
    "carboidrati_totali_100g",
    "carboidrati_totali_totale",
    "di_cui_zuccheri_100g",
    "di_cui_zuccheri_totale",
    "grassi_totali_100g",
    "grassi_totali_totale",
    "di_cui_saturi_100g",
    "di_cui_saturi_totale",
    "monoinsaturi_100g",
    "monoinsaturi_totale",
    "polinsaturi_100g",
    "polinsaturi_totale",
    "proteine_totali_100g",
    "proteine_totali_totale",
    "colesterolo_totale_100g",
    "colesterolo_totale_totale",
    "fibre_100g",
    "fibre_totale",
    "sodio_100g",
    "sodio_totale",
    "spesa_totale_acquisto",
    "spesa_totale_ricetta",
    "spesa_per_porzione",
]

_TEMPLATE_REQUIRED_LIST_FIELDS = [
    "allergeni_loghi",
    "allergeni_tracce_loghi",
    "ingredienti_dettaglio",
]

_CLOUD_REQUIRED_KEYS = [
    "titolo",
    "ingredienti_blocco",
    "porzioni",
    "tempo_dettaglio",
    "difficolta",
    "procedimento_blocco",
    "conservazione",
    "categoria",
    "vegetariano flag",
    "diete",
    "note errori",
    "vino descrizione",
    "vino temperatura servizio",
    "vino regione",
    "vino annata",
    "vino motivo annata",
    "allergeni elenco",
    "attrezzature specifiche",
    "attrezzature generiche",
    "attrezzature semplici",
    "attrezzature professionali",
    "attrezzature pasticceria",
    "presentazione impiattamento",
    "ingredienti_dettaglio",
    "stagionalita",
    "spesa totale acquisto",
    "spesa totale ricetta",
    "spesa per porzione",
    "energia 100g",
    "energia totale",
    "carboidrati totali 100g",
    "carboidrati totali totale",
    "di cui zuccheri 100g",
    "di cui zuccheri totale",
    "grassi totali 100g",
    "grassi totali totale",
    "di cui saturi 100g",
    "di cui saturi totale",
    "monoinsaturi 100g",
    "monoinsaturi totale",
    "polinsaturi 100g",
    "polinsaturi totale",
    "proteine totali 100g",
    "proteine totali totale",
    "colesterolo totale 100g",
    "colesterolo totale totale",
    "fibre 100g",
    "fibre totale",
    "sodio 100g",
    "sodio totale",
]


def _collect_missing_template_fields(ctx: Dict[str, Any]) -> List[str]:
    fields = set()

    def _missing_text(key: str) -> None:
        val = ctx.get(key)
        if isinstance(val, bool):
            return
        if _is_empty_value(val):
            fields.add(key)

    def _missing_num(key: str) -> None:
        val = ctx.get(key)
        if _is_empty_value(val) or _is_zero_like(val):
            fields.add(key)

    for key in _TEMPLATE_REQUIRED_TEXT_FIELDS:
        _missing_text(key)

    for key in _TEMPLATE_REQUIRED_NUM_FIELDS:
        _missing_num(key)

    for key in _TEMPLATE_REQUIRED_LIST_FIELDS:
        val = ctx.get(key)
        if _is_empty_value(val):
            fields.add(key)

    rows = ctx.get("ingredienti_dettaglio")
    if isinstance(rows, list) and rows:
        required_row_keys = [
            "ingrediente",
            "scarto",
            "peso_min_acquisto",
            "prezzo_kg_ud",
            "quantita_usata",
            "prezzo_alimento_acquisto",
            "prezzo_calcolato",
        ]
        for idx, row in enumerate(rows):
            if not isinstance(row, dict):
                fields.add(f"ingredienti_dettaglio[{idx}]")
                continue
            for rkey in required_row_keys:
                if _is_empty_value(row.get(rkey)):
                    fields.add(f"ingredienti_dettaglio.{rkey}")

    return sorted(fields)


def _merge_missing_fields(dst: Dict[str, Any], src: Dict[str, Any]) -> List[str]:
    applied: List[str] = []
    if not isinstance(src, dict):
        return applied

    def set_if_empty(key: str, value: Any, zero_missing: bool = False) -> None:
        if not _is_empty_value(value) and (
            _is_empty_value(dst.get(key)) or (zero_missing and _is_zero_like(dst.get(key)))
        ):
            dst[key] = value
            applied.append(key)

    def _as_text(value: Any) -> Any:
        if isinstance(value, list):
            return ", ".join([str(x).strip() for x in value if str(x).strip()])
        return value

    def _merge_diet_text(existing: Any, incoming: Any) -> str:
        def _split(val: Any) -> List[str]:
            if val is None:
                return []
            s = str(val)
            if not s.strip():
                return []
            parts = re.split(r"[;,\n]+", s)
            out: List[str] = []
            for p in parts:
                p = p.strip().strip("-").strip()
                if p:
                    out.append(p)
            return out

        items = _split(existing) + _split(incoming)
        seen: set[str] = set()
        merged: List[str] = []
        for item in items:
            key = item.lower()
            if key in seen:
                continue
            seen.add(key)
            merged.append(item)
        return ", ".join(merged).strip()

    for key in (
        "title",
        "category",
        "servings",
        "difficulty",
        "prep_time_min",
        "cook_time_min",
        "total_time_min",
        "diet_text",
        "conservazione",
        "allergens_text",
        "vino_descrizione",
        "wine_pairing",
        "vino_temperatura_servizio",
        "vino_regione",
        "vino_annata",
        "vino_motivo_annata",
        "equipment_text",
        "attrezzature_generiche",
        "attrezzature_semplici",
        "attrezzature_professionali",
        "attrezzature_pasticceria",
        "presentazione_impiattamento",
        "stagionalita",
        "notes",
    ):
        if key in src:
            val = src.get(key)
            if key in {
                "allergens_text",
                "equipment_text",
                "attrezzature_generiche",
                "attrezzature_semplici",
                "attrezzature_professionali",
                "attrezzature_pasticceria",
                "vino_descrizione",
                "wine_pairing",
                "vino_temperatura_servizio",
                "vino_regione",
                "vino_annata",
                "vino_motivo_annata",
                "conservazione",
                "presentazione_impiattamento",
                "diet_text",
                "stagionalita",
            }:
                val = _as_text(val)
            if key == "diet_text":
                src_text = _as_text(val)
                dst_text = dst.get(key)
                if not _is_empty_value(src_text) and not _is_empty_value(dst_text):
                    merged = _merge_diet_text(dst_text, src_text)
                    if merged and merged != str(dst_text):
                        dst[key] = merged
                        applied.append(key)
                        continue
            set_if_empty(key, val)

    for key in ("spesa_totale_ricetta", "spesa_per_porzione", "spesa_totale_acquisto"):
        if key in src:
            set_if_empty(key, src.get(key), zero_missing=True)

    if "fonte_prezzi" in src:
        set_if_empty("fonte_prezzi", src.get("fonte_prezzi"))

    src_ingredients = src.get("ingredients")
    if _is_empty_value(dst.get("ingredients")) and isinstance(src_ingredients, list) and src_ingredients:
        dst["ingredients"] = src_ingredients
        applied.append("ingredients")

    src_steps = src.get("steps")
    if _is_empty_value(dst.get("steps")) and isinstance(src_steps, list) and src_steps:
        dst["steps"] = src_steps
        applied.append("steps")

    src_allergens = src.get("allergens")
    if not src_allergens:
        src_allergens = src.get("allergens_present")
    if _is_empty_value(dst.get("allergens")) and isinstance(src_allergens, list) and src_allergens:
        dst["allergens"] = src_allergens
        applied.append("allergens")

    src_traces = src.get("allergens_traces")
    if not src_traces:
        src_traces = src.get("traces_allergens")
    if _is_empty_value(dst.get("traces_allergens")) and isinstance(src_traces, list) and src_traces:
        dst["traces_allergens"] = src_traces
        applied.append("traces_allergens")

    src_nt = src.get("nutrition_table")
    if isinstance(src_nt, dict) and src_nt:
        dst_nt = dst.get("nutrition_table")
        if not isinstance(dst_nt, dict):
            dst["nutrition_table"] = src_nt
            applied.append("nutrition_table")
        else:
            updated = False
            for scope in ("100g", "totale", "porzione"):
                src_block = src_nt.get(scope)
                if not isinstance(src_block, dict):
                    continue
                dst_block = dst_nt.get(scope)
                if not isinstance(dst_block, dict):
                    dst_block = {}
                    dst_nt[scope] = dst_block
                for key, val in src_block.items():
                    if _is_empty_value(val) or _is_zero_like(val):
                        continue
                    if _is_empty_value(dst_block.get(key)) or _is_zero_like(dst_block.get(key)):
                        dst_block[key] = val
                        updated = True
            if updated:
                dst["nutrition_table"] = dst_nt
                applied.append("nutrition_table")

    src_cost = src.get("cost_lines")
    if isinstance(src_cost, list) and src_cost:
        dst_cost = dst.get("cost_lines")
        if not isinstance(dst_cost, list) or not dst_cost:
            dst["cost_lines"] = src_cost
            applied.append("cost_lines")
        else:
            def _norm_name(value: Any) -> str:
                return re.sub(r"\s+", " ", str(value or "").strip().lower())

            name_to_idx: Dict[str, int] = {}
            for idx, row in enumerate(dst_cost):
                if not isinstance(row, dict):
                    continue
                key = _norm_name(row.get("ingrediente") or row.get("ingredient") or row.get("name"))
                if key and key not in name_to_idx:
                    name_to_idx[key] = idx

            updated = False
            for src_row in src_cost:
                if not isinstance(src_row, dict):
                    continue
                src_name = _norm_name(src_row.get("ingrediente") or src_row.get("ingredient") or src_row.get("name"))
                match_idx = name_to_idx.get(src_name) if src_name else None
                if match_idx is None:
                    for idx, row in enumerate(dst_cost):
                        if isinstance(row, dict) and not _norm_name(row.get("ingrediente") or row.get("ingredient") or row.get("name")):
                            match_idx = idx
                            break
                if match_idx is None:
                    dst_cost.append(dict(src_row))
                    updated = True
                    continue
                dst_row = dst_cost[match_idx]
                if not isinstance(dst_row, dict):
                    dst_cost[match_idx] = dict(src_row)
                    updated = True
                    continue
                scarto_keys = {"scarto", "scarto_pct", "waste_pct"}
                for key, val in src_row.items():
                    if _is_empty_value(val):
                        continue
                    if key not in scarto_keys and _is_zero_like(val):
                        continue
                    if _is_empty_value(dst_row.get(key)) or (
                        key not in scarto_keys and _is_zero_like(dst_row.get(key))
                    ):
                        dst_row[key] = val
                        updated = True
            if updated:
                dst["cost_lines"] = dst_cost
                applied.append("cost_lines")

    src_flags = src.get("diet_flags")
    if isinstance(src_flags, dict):
        dst_flags = dst.get("diet_flags")
        has_dst = isinstance(dst_flags, dict) and any(bool(v) for v in dst_flags.values())
        if not has_dst and any(bool(v) for v in src_flags.values()):
            dst["diet_flags"] = {
                "vegetarian": bool(src_flags.get("vegetarian")),
                "vegan": bool(src_flags.get("vegan")),
                "gluten_free": bool(src_flags.get("gluten_free")),
                "lactose_free": bool(src_flags.get("lactose_free")),
            }
            applied.append("diet_flags")

    if dst.get("vino_descrizione") and not dst.get("wine_pairing"):
        dst["wine_pairing"] = dst.get("vino_descrizione")
    if dst.get("wine_pairing") and not dst.get("vino_descrizione"):
        dst["vino_descrizione"] = dst.get("wine_pairing")
    if dst.get("equipment_text") and not dst.get("attrezzature_generiche"):
        dst["attrezzature_generiche"] = dst.get("equipment_text")
    if dst.get("attrezzature_generiche") and not dst.get("equipment_text"):
        dst["equipment_text"] = dst.get("attrezzature_generiche")

    return applied


def _apply_ai_patch(
    recipe: Dict[str, Any],
    ai_recipe: Dict[str, Any],
    *,
    allow_override_lists: bool = False,
) -> List[str]:
    applied = _merge_missing_fields(recipe, ai_recipe)
    if not isinstance(ai_recipe, dict):
        return applied

    if allow_override_lists:
        ings = ai_recipe.get("ingredients")
        if isinstance(ings, list):
            cleaned_ings: List[Dict[str, Any]] = []
            for ing in ings:
                if isinstance(ing, dict):
                    name = str(ing.get("name") or "").strip()
                    if not name:
                        continue
                    cleaned_ings.append(
                        {
                            "name": name,
                            "qty": ing.get("qty"),
                            "unit": ing.get("unit"),
                        }
                    )
                else:
                    name = str(ing or "").strip()
                    if not name:
                        continue
                    cleaned_ings.append({"name": name, "qty": None, "unit": None})
            if cleaned_ings:
                recipe["ingredients"] = cleaned_ings
                applied.append("ingredients_normalized")

        steps = ai_recipe.get("steps")
        if isinstance(steps, list):
            cleaned_steps: List[Dict[str, Any]] = []
            for st in steps:
                if isinstance(st, dict):
                    txt = _clean_step_text(st.get("text") or st.get("step") or "")
                else:
                    txt = _clean_step_text(st)
                if not txt:
                    continue
                cleaned_steps.append({"text": txt})
            if cleaned_steps:
                recipe["steps"] = cleaned_steps
                applied.append("steps_normalized")

    return applied


def _ollama_complete_recipe(
    recipe: Dict[str, Any],
    source_text: str,
    missing_fields: List[str],
) -> str:
    src = (source_text or "").strip()
    if len(src) > MAX_SOURCE_CHARS:
        src = src[:MAX_SOURCE_CHARS] + "\n...[TRONCATO]"

    schema = {
        "title": "string|null",
        "category": "string|null",
        "servings": 0,
        "difficulty": "bassa|media|alta",
        "prep_time_min": 0,
        "cook_time_min": 0,
        "total_time_min": 0,
        "diet_text": "string|null",
        "diet_flags": {"vegetarian": False, "vegan": False, "gluten_free": False, "lactose_free": False},
        "conservazione": "string|null",
        "allergens_text": "string|null",
        "allergens_present": ["glutine", "crostacei", "uova", "pesce", "arachidi", "soia", "latte", "frutta_a_guscio", "sedano", "senape", "sesamo", "solfiti", "lupini", "molluschi"],
        "allergens_traces": ["glutine", "crostacei", "uova", "pesce", "arachidi", "soia", "latte", "frutta_a_guscio", "sedano", "senape", "sesamo", "solfiti", "lupini", "molluschi"],
        "wine_pairing": "string|null",
        "equipment_text": "string|null",
        "equipment_simple": ["string"],
        "equipment_professional": ["string"],
        "equipment_pasticceria": ["string"],
        "presentazione_impiattamento": "string|null",
        "ingredients": [{"name": "string", "qty": 0, "unit": "g|kg|ml|l|pz|cucchiai|cucchiaini|q.b."}],
        "steps": ["string"],
        "nutrition_table": {
            "100g": {"energia": 0, "carboidrati_totali": 0, "di_cui_zuccheri": 0, "grassi_totali": 0, "di_cui_saturi": 0, "proteine_totali": 0, "fibre": 0, "sodio": 0},
            "totale": {"energia": 0, "carboidrati_totali": 0, "di_cui_zuccheri": 0, "grassi_totali": 0, "di_cui_saturi": 0, "proteine_totali": 0, "fibre": 0, "sodio": 0},
            "porzione": {"energia": 0, "carboidrati_totali": 0, "di_cui_zuccheri": 0, "grassi_totali": 0, "di_cui_saturi": 0, "proteine_totali": 0, "fibre": 0, "sodio": 0}
        },
        "cost_lines": [
            {
                "ingrediente": "string",
                "scarto": "string|null",
                "peso_min_acquisto": "string|null",
                "prezzo_kg_ud": "string|null",
                "quantita_usata": "string|null",
                "prezzo_alimento_acquisto": "string|null",
                "prezzo_calcolato": "string|null"
            }
        ],
        "spesa_totale_acquisto": "string|null",
        "spesa_totale_ricetta": "string|null",
        "spesa_per_porzione": "string|null",
        "fonte_prezzi": "string|null",
    }

    prompt = (
        "Sei un assistente chef. Completa una ricetta italiana.\n"
        "Se Ingredienti/Procedimento non sono presenti nel testo, crea una versione plausibile coerente col titolo.\n"
        "Se mancano costi o valori nutrizionali, stima valori plausibili.\n"
        "Compila 'diet_text' con le diete compatibili, scegliendo SOLO tra: Dieta vegetariana, Dieta vegana, "
        "Dieta plant-based, Dieta flexitariana, Dieta pescetariana, Dieta senza glutine (per celiachia), "
        "Dieta ipocalorica, Dieta ipercalorica, Dieta per diabetici, Dieta per ipertensione, "
        "Dieta per reflusso gastroesofageo, Dieta halal, Dieta kosher, Dieta induista, Dieta buddhista, "
        "Dieta mediterranea, Dieta chetogenica.\n"
        "Rispondi SOLO con un JSON valido secondo lo schema.\n"
        f"SCHEMA: {json.dumps(schema, ensure_ascii=False)}\n"
        f"CAMPI MANCANTI: {json.dumps(missing_fields, ensure_ascii=False)}\n"
        f"RICETTA ESISTENTE: {json.dumps(recipe, ensure_ascii=False)}\n"
        f"TESTO ORIGINALE:\n{src}\n"
    )
    return _ask_ollama_json(prompt)


def _complete_missing_with_ai(
    recipe: Dict[str, Any],
    source_text: str,
    missing_fields: List[str],
    *,
    allow_cloud: bool = True,
    allow_ollama: bool = False,
    retries_cloud: int = 2,
    retries_ollama: int = 1,
    subscription_tier: str = "",
    allow_shared_cloud: bool = False,
) -> Tuple[Optional[Dict[str, Any]], Optional[str], Optional[str]]:
    if not missing_fields:
        return None, None, None

    error_msg: Optional[str] = None
    provider_used: Optional[str] = None

    if allow_cloud and cloud_ai_settings is not None:
        for _ in range(max(1, int(retries_cloud))):
            try:
                used, patch, provider = cloud_ai_settings.complete_missing_fields(
                    recipe,
                    missing_fields,
                    source_text,
                    subscription_tier=subscription_tier,
                    allow_shared=allow_shared_cloud,
                )
                provider_used = provider
                if used and isinstance(patch, dict) and patch:
                    ai_recipe, _ = parse_recipe_text(json.dumps(patch, ensure_ascii=False))
                    return ai_recipe, provider, None
                if used and not patch:
                    error_msg = f"cloud_ai: risposta vuota ({provider})"
            except Exception as e:
                error_msg = f"cloud_ai: {e}"

    if allow_cloud and ai_cloud is not None:
        for _ in range(max(1, int(retries_cloud))):
            try:
                patch, meta = ai_cloud.request_full_recipe(recipe, source_text, missing_fields)
                if isinstance(patch, dict) and patch:
                    ai_recipe, _ = parse_recipe_text(json.dumps(patch, ensure_ascii=False))
                    provider = meta.get("provider") if isinstance(meta, dict) else None
                    return ai_recipe, str(provider) if provider else "cloud", None
                if patch is not None and not patch:
                    error_msg = "cloud_ai: risposta vuota (legacy)"
            except Exception as e:
                error_msg = f"cloud_ai: {e}"

    if allow_cloud and cloud_ai_settings is None and ai_cloud is None and error_msg is None:
        error_msg = "cloud_ai: provider non configurato"

    if allow_ollama:
        for _ in range(max(1, int(retries_ollama))):
            try:
                ai_text = _ollama_complete_recipe(recipe, source_text, missing_fields)
                if ai_text:
                    ai_recipe, _ = parse_recipe_text(ai_text)
                    if isinstance(ai_recipe, dict) and ai_recipe:
                        return ai_recipe, "ollama", None
            except Exception:
                error_msg = "ollama: errore richiesta"

    return None, provider_used, error_msg


def _cloud_ai_available(*, allow_shared: bool = False, subscription_tier: str = "") -> bool:
    tier = str(subscription_tier or "").lower().strip()
    tier_allows_shared = tier not in {"", "free", "starter"}
    if cloud_ai_settings is not None:
        try:
            s = cloud_ai_settings.load_settings()
            if s.get("enabled"):
                prov = str(s.get("provider") or "auto").lower().strip()
                has_openai = bool((s.get("openai") or {}).get("api_key"))
                has_gemini = bool((s.get("gemini") or {}).get("api_key"))
                if prov == "openai":
                    return has_openai
                if prov == "gemini":
                    return has_gemini
                if prov == "auto":
                    return has_openai or has_gemini
            if allow_shared and tier_allows_shared:
                if getattr(cloud_ai_settings, "shared_openai_available", None):
                    if cloud_ai_settings.shared_openai_available():
                        return True
        except Exception:
            pass

    if ai_cloud is not None:
        try:
            return ai_cloud.pick_provider() is not None
        except Exception:
            pass

    return False


def estimate_nutrition_ai(ingredients_text: str) -> Dict[str, Any]:
    prompt = (
        "Stima i valori per questa ricetta in formato JSON rigoroso:\n"
        f"Ingredienti: {ingredients_text}\n"
        "Rispondi SOLO con questo JSON:\n"
        '{"cost_eur": 0.0, "kcal": 0, "carbs": 0, "fats": 0, "proteins": 0}'
    )
    resp = _ask_ollama_json(prompt)
    try:
        if "```" in resp:
            if "json" in resp:
                parts = resp.split("```json")
                resp = parts[-1].split("```")[0] if len(parts) > 1 else resp
            else:
                parts = resp.split("```")
                resp = parts[1] if len(parts) > 1 else resp
        return json.loads(resp)
    except Exception:
        return {}


# ------------------------------------------------------------
# Business rules and enrichment
# ------------------------------------------------------------
def _fmt_qty(qty: Any) -> str:
    try:
        q = float(qty)
    except Exception:
        return ""
    if abs(q) < 1e-9:
        return ""
    if q.is_integer():
        return str(int(q))
    return str(q).rstrip("0").rstrip(".")


def _build_ingredients_text(ingredients: List[Dict[str, Any]]) -> str:
    lines: List[str] = []
    for ing in ingredients or []:
        if not isinstance(ing, dict):
            continue
        name = str(ing.get("name") or "").strip()
        if not name:
            continue
        qty = _fmt_qty(ing.get("qty"))
        unit = str(ing.get("unit") or "").strip()
        parts = [p for p in [qty, unit, name] if p]
        line = " ".join(parts).strip()
        if line:
            lines.append(f"- {line}")
    return "\n".join(lines).strip()


def _clean_step_text(text: Any) -> str:
    s = str(text or "").strip()
    if not s:
        return ""
    s = re.sub(r"^\s*[\-\*\u2022]+\s*", "", s)
    s = re.sub(r"^\s*\d+(?:[\.\)\-]+|\s)+", "", s)
    s = s.strip()
    if not s or re.fullmatch(r"\d+(?:[\.\)\-]+)?", s):
        return ""
    return s


def _build_steps_text(steps: List[Dict[str, Any]], *, numbered: bool = True) -> str:
    lines: List[str] = []
    for idx, st in enumerate(steps or []):
        if isinstance(st, dict):
            txt = _clean_step_text(st.get("text"))
        else:
            txt = _clean_step_text(st)
        if not txt:
            continue
        if numbered:
            lines.append(f"{idx + 1}. {txt}")
        else:
            lines.append(txt)
    return "\n".join(lines).strip()


def _title_from_paths(paths: List[str]) -> str:
    if not paths:
        return "Ricetta"
    stem = Path(paths[0]).stem
    stem = stem.replace("_", " ").replace("-", " ")
    stem = re.sub(r"\s+", " ", stem).strip()
    return stem or "Ricetta"


def _apply_saverio_rules(recipe: Dict[str, Any]) -> None:
    title_raw = recipe.get("title") or recipe.get("titolo")
    if isinstance(title_raw, str) and title_raw.strip():
        t = re.sub(r"\bTitolo\s*:\s*", "", title_raw, flags=re.IGNORECASE).strip()
        if "|" in t:
            parts = [p.strip() for p in t.split("|") if p.strip()]
            uniq: List[str] = []
            for p in parts:
                if p.lower() not in [u.lower() for u in uniq]:
                    uniq.append(p)
            t = uniq[0] if uniq else t
        recipe["title"] = t

    if not recipe.get("difficulty"):
        diff_raw = recipe.get("difficolta") or recipe.get("difficoltà")
        if isinstance(diff_raw, str) and diff_raw.strip():
            recipe["difficulty"] = diff_raw.strip().lower()

    if recipe.get("servings") is None:
        serv_raw = recipe.get("porzioni")
        if serv_raw is not None:
            try:
                recipe["servings"] = int(float(str(serv_raw).replace(",", ".")))
            except Exception:
                m = re.search(r"(\d+)", str(serv_raw))
                if m:
                    recipe["servings"] = int(m.group(1))

    if recipe.get("prep_time_min") is None or recipe.get("cook_time_min") is None:
        tempo_raw = str(recipe.get("tempo_dettaglio") or recipe.get("tempo") or "").strip()

        def _parse_minutes(raw: str) -> Optional[int]:
            s = (raw or "").strip().lower()
            if not s:
                return None
            h = 0
            m = 0
            mh = re.search(r"(\d+)\s*h", s)
            if mh:
                try:
                    h = int(mh.group(1))
                except Exception:
                    h = 0
            mm = re.search(r"(\d+)\s*(min|m)\b", s)
            if mm:
                try:
                    m = int(mm.group(1))
                except Exception:
                    m = 0
            if h or m:
                return h * 60 + m
            mn = re.search(r"(\d+)", s)
            if mn:
                try:
                    return int(mn.group(1))
                except Exception:
                    return None
            return None

        def _extract_time(label: str, raw: str) -> Optional[int]:
            m = re.search(rf"{label}\s*[:\-]?\s*([0-9][^,;]*)", (raw or "").lower())
            if not m:
                return None
            return _parse_minutes(m.group(1))

        if tempo_raw:
            prep = _extract_time(r"(?:prep|preparazione)", tempo_raw)
            cook = _extract_time(r"(?:cottura|cook)", tempo_raw)
            total = _extract_time(r"(?:totale|complessivo|total)", tempo_raw)
            if recipe.get("prep_time_min") is None and prep is not None:
                recipe["prep_time_min"] = prep
            if recipe.get("cook_time_min") is None and cook is not None:
                recipe["cook_time_min"] = cook
            if recipe.get("total_time_min") is None and total is not None:
                recipe["total_time_min"] = total

        if recipe.get("total_time_min") is None:
            prep_val = recipe.get("prep_time_min")
            cook_val = recipe.get("cook_time_min")
            try:
                total = int((int(prep_val) if prep_val else 0) + (int(cook_val) if cook_val else 0))
            except Exception:
                total = 0
            if total > 0:
                recipe["total_time_min"] = total

    if isinstance(recipe.get("ingredients"), list):
        clean_list = []
        for ing in recipe["ingredients"]:
            if not isinstance(ing, dict):
                continue

            name = str(ing.get("name") or "").strip()
            if not name:
                continue

            unit = ing.get("unit")
            if isinstance(unit, str):
                ing["unit"] = unit.strip()

            ing["name"] = name
            clean_list.append(ing)
        recipe["ingredients"] = clean_list

    if isinstance(recipe.get("steps"), list):
        clean_steps = []
        for st in recipe["steps"]:
            if isinstance(st, dict):
                txt = str(st.get("text") or "").strip()
                if not txt:
                    continue
                st["text"] = txt
                clean_steps.append(st)
            else:
                txt = str(st).strip()
                if not txt:
                    continue
                clean_steps.append({"text": txt})
        recipe["steps"] = clean_steps

    conservazione = recipe.get("conservazione")
    if isinstance(conservazione, str) and conservazione.strip():
        title_l = str(recipe.get("title") or "").lower()
        cons = conservazione.strip()
        if "risotto" in cons.lower() and "risotto" not in title_l:
            parts = re.split(r"(?<=[\.\!\?])\s+", cons)
            parts = [p for p in parts if "risotto" not in p.lower()]
            cons = " ".join(parts).strip()
        cons = re.sub(r"\bvalori descritti per grammi\b.*", "", cons, flags=re.IGNORECASE).strip()
        recipe["conservazione"] = cons

    title_l = str(recipe.get("title") or "").lower()
    force_molecular = "molecolare" in title_l

    valid_cats = [
        "Antipasti",
        "Primi",
        "Secondi",
        "Contorni",
        "Dolci",
        "Pane e lievitati",
        "Salse e condimenti",
        "Bevande",
    ]
    cat_raw = str(recipe.get("category") or recipe.get("categoria") or "Altro").strip()
    cat_lower = cat_raw.lower()
    found = False
    if force_molecular:
        recipe["category"] = "Molecolare"
        found = True
    elif "antipasto" in cat_lower:
        recipe["category"] = "Antipasti"
        found = True
    elif "primo" in cat_lower:
        recipe["category"] = "Primi"
        found = True
    elif "secondo" in cat_lower:
        recipe["category"] = "Secondi"
        found = True
    elif "dessert" in cat_lower:
        recipe["category"] = "Dolci"
        found = True
    elif "piatto unico" in cat_lower:
        recipe["category"] = "Piatto unico"
        found = True
    else:
        cat = cat_raw.capitalize()
        for vc in valid_cats:
            if vc.lower() in cat.lower():
                recipe["category"] = vc
                found = True
                break
    if not found:
        recipe["category"] = "Altro"


def _nutrition_table_from(nutrition: Dict[str, Any]) -> Dict[str, Dict[str, float]]:
    table: Dict[str, Dict[str, float]] = {"100g": {}, "totale": {}, "porzione": {}}
    if not nutrition:
        return table

    total = nutrition.get("total") if isinstance(nutrition, dict) else None
    per_100g = nutrition.get("per_100g") if isinstance(nutrition, dict) else None
    per_portion = nutrition.get("per_portion") if isinstance(nutrition, dict) else None

    def pick(src: Dict[str, Any], key: str) -> Optional[float]:
        try:
            val = src.get(key)
            if val is None:
                return None
            return float(val)
        except Exception:
            return None

    def fill(dst: Dict[str, float], src: Optional[Dict[str, Any]]) -> None:
        if not isinstance(src, dict):
            return
        kcal = pick(src, "kcal")
        carbs = pick(src, "carbs_g")
        sugars = pick(src, "sugars_g")
        fats = pick(src, "fat_g") or pick(src, "fats_g")
        sat = pick(src, "sat_fat_g")
        proteins = pick(src, "protein_g") or pick(src, "proteins_g")
        fibre = pick(src, "fibre_g") or pick(src, "fiber_g")
        salt_g = pick(src, "salt_g")

        if kcal is not None:
            dst["energia"] = kcal
        if carbs is not None:
            dst["carboidrati_totali"] = carbs
        if sugars is not None:
            dst["di_cui_zuccheri"] = sugars
        if fats is not None:
            dst["grassi_totali"] = fats
        if sat is not None:
            dst["di_cui_saturi"] = sat
        if proteins is not None:
            dst["proteine_totali"] = proteins
        if fibre is not None:
            dst["fibre"] = fibre
        if salt_g is not None:
            dst["sodio"] = salt_g * 1000.0

    fill(table["totale"], total)
    fill(table["100g"], per_100g)
    fill(table["porzione"], per_portion)
    return table


def _cost_lines_from_summary(cost_summary: Dict[str, Any]) -> List[Dict[str, Any]]:
    out: List[Dict[str, Any]] = []
    if not isinstance(cost_summary, dict):
        return out
    items = cost_summary.get("items")
    if not isinstance(items, list):
        return out

    def _normalize_unit_label(value: Any) -> str:
        s = str(value or "").strip()
        if not s:
            return ""
        return re.sub(r"\b(ud|u|unita|unità|unit)\b", "pz", s, flags=re.IGNORECASE)

    def _normalize_price_unit(value: Any) -> str:
        s = str(value or "").strip()
        if not s:
            return ""
        s = s.replace("?/","€/")
        if "/" in s and "€" not in s and "eur" not in s.lower():
            s = s.replace("/", "€/", 1)
        return s

    for it in items:
        if not isinstance(it, dict):
            continue
        raw = str(it.get("ingredient_raw") or it.get("ingredient") or "").strip()
        if not raw:
            continue
        qty = it.get("qty")
        unit = _normalize_unit_label(it.get("unit_raw"))
        quantita = ""
        if qty is not None and unit:
            quantita = f"{qty} {unit}"
        elif qty is not None:
            quantita = str(qty)

        price_unit = _normalize_price_unit(it.get("price_unit") or "")
        price_val = it.get("price_value")
        prezzo_kg = ""
        if price_val is not None and price_unit:
            prezzo_kg = f"{price_val} {price_unit}"

        cost_val = it.get("cost")
        prezzo_calcolato = f"{cost_val:.2f}" if isinstance(cost_val, (int, float)) else ""
        scarto_val = it.get("waste_pct") or it.get("scarto") or it.get("scarto_pct")
        scarto = ""
        if scarto_val is not None:
            try:
                scarto_num = float(scarto_val)
                scarto = "" if scarto_num == 0 else f"{scarto_num:g}"
            except Exception:
                scarto = str(scarto_val).strip()

        out.append(
            {
                "ingrediente": raw,
                "peso_min_acquisto": "",
                "prezzo_kg_ud": prezzo_kg,
                "quantita_usata": quantita,
                "prezzo_alimento_acquisto": "",
                "prezzo_calcolato": prezzo_calcolato,
                "scarto": scarto,
            }
        )
    return out


def _estimate_waste_pct(name: Any) -> Optional[str]:
    s = str(name or "").lower()
    s = re.sub(r"[^a-z0-9\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        return None

    def has_any(keys: List[str]) -> bool:
        for kw in keys:
            if " " in kw:
                if kw in s:
                    return True
            else:
                if re.search(r"\b" + re.escape(kw), s):
                    return True
        return False

    if has_any(["al netto", "gia pulito", "già pulito", "pulito", "sbucciato", "pelato", "filetto", "filetti", "disossato", "senza pelle", "senza lische"]):
        return "5"

    if has_any([
        "surgelat", "pronto", "precotto", "conserva", "in scatola", "in barattolo",
        "passata", "pelati", "purea", "polpa di pomodoro", "concentrato di pomodoro",
        "latte", "burro", "panna", "yogurt", "formagg", "ricotta", "mascarpone",
        "mozzarella", "parmig", "grana", "pecorino", "gorgonzola",
        "uovo", "uova", "olio", "acqua", "sale", "zucchero", "farina", "cacao",
        "cioccolato", "lievito", "vanillina", "riso", "pasta", "pane", "miele",
        "aceto", "vino", "birra", "spezie", "pepe", "brodo", "amido", "fecola",
        "semola", "gelatina", "caffe", "the", "te",
        "prosciutto", "speck", "salame", "pancetta", "bresaola", "wurstel",
    ]):
        return "0"

    if has_any(["filetto", "filetti", "trancio", "carpaccio", "petto", "fesa", "lombata", "macinat", "bistecc", "scalopp", "hamburger"]):
        return "5"

    if has_any(["pollo intero", "gallina intera", "tacchino intero", "coniglio intero", "anatra intera", "faraona", "carcassa", "ali", "coscia", "sovracoscia", "costine", "costole"]):
        return "25"

    if ("intero" in s or "intera" in s) and has_any(["pesce", "orata", "branzino", "spigola", "salmone", "tonno", "merluzz", "sogliola", "nasello", "trota", "baccala", "sgombro", "sardina", "alici", "acciug", "cefalo", "rombo"]):
        return "40"

    if has_any(["gamber", "scampi", "crostace", "cozza", "cozze", "vongol", "mollusch", "ostrica", "capesant", "canestrel", "aragost", "astice"]):
        return "55"

    if has_any(["calamar", "seppia", "totano"]):
        return "25"
    if has_any(["polpo"]):
        return "15"

    if has_any(["pesce", "orata", "branzino", "spigola", "salmone", "tonno", "merluzz", "sogliola", "nasello", "trota", "baccala", "sgombro", "sardina", "alici", "acciug", "cefalo", "rombo"]):
        return "35"

    if has_any(["carne", "manzo", "vitello", "maiale", "pollo", "tacchino", "agnello", "bov", "suin", "coniglio", "cervo", "salsicc"]):
        return "8"

    if has_any(["tubero", "patata", "patate", "batata", "topinambur", "manioca"]):
        return "5"

    if has_any(["carciof", "asparag", "cardo"]):
        return "40"
    if has_any(["cavolfior", "broccol", "cavol", "verza", "cappucc", "zucca"]):
        return "30"
    if has_any(["finocch", "porro", "sedan", "rapa"]):
        return "25"
    if has_any(["insalata", "lattuga", "spinac", "bietol", "rucola", "basilico", "prezzemolo", "rosmarino", "salvia", "timo", "menta", "coriandolo", "erba cipollina", "erbette"]):
        return "20"
    if has_any(["patat", "carot", "cipoll", "aglio", "zucchin", "melanzan", "peperon", "pomodor", "cetriol", "ravanell"]):
        return "12"
    if has_any(["fung"]):
        return "5"

    if has_any(["ananas"]):
        return "45"
    if has_any(["banana", "agrum", "arancia", "limone", "mandarino", "pompelmo", "melone", "anguria", "avocado", "mango", "papaya"]):
        return "35"
    if has_any(["kiwi"]):
        return "20"
    if has_any(["mela", "pera", "pesca", "albicocc", "susin", "prugn", "cilieg", "fico"]):
        return "15"
    if has_any(["uva", "mirtill", "lampon", "fragol", "frutti di bosco"]):
        return "5"

    if "guscio" in s and has_any(["mandorl", "nocciol", "noci", "pistacch", "arachid", "anacard", "pinoli"]):
        return "45"
    if has_any(["mandorl", "nocciol", "noci", "pistacch", "arachid", "anacard", "pinoli"]):
        return "5"

    return "5"


def _coerce_str_dict(obj: Any) -> Dict[str, Any]:
    data: Any = None
    if isinstance(obj, dict):
        data = obj
    elif hasattr(obj, "to_dict"):
        try:
            data = obj.to_dict()
        except Exception:
            data = None
    if data is None:
        try:
            data = dict(getattr(obj, "__dict__", {}) or {})
        except Exception:
            data = None
    if not isinstance(data, dict):
        return {}
    out: Dict[str, Any] = {}
    for k, v in data.items():
        out[str(k)] = v
    return out


def _enrich_data(
    recipe: Dict[str, Any],
    source_text: str = "",
    *,
    allow_local_ai: bool = True,
) -> None:
    base_dir = Path(__file__).resolve().parent.parent

    cost_summary: Dict[str, Any] = {}
    nutrition_dict: Dict[str, Any] = {}
    cost_lines_raw = recipe.get("cost_lines")
    has_cost_lines = isinstance(cost_lines_raw, list) and bool(cost_lines_raw)
    has_cost_values = False
    if isinstance(cost_lines_raw, list):
        for row in cost_lines_raw:
            if not isinstance(row, dict):
                continue
            for key in (
                "prezzo_calcolato",
                "prezzo_kg_ud",
                "prezzo_alimento_acquisto",
                "prezzo_acquisto",
                "price_value",
                "cost",
                "scarto",
                "scarto_pct",
                "waste_pct",
            ):
                if not _is_empty_value(row.get(key)) and not _is_zero_like(row.get(key)):
                    has_cost_values = True
                    break
            if has_cost_values:
                break
    if has_cost_lines and not has_cost_values:
        has_cost_lines = False
    has_cost_tot = not _is_empty_value(recipe.get("spesa_totale_ricetta")) and not _is_zero_like(recipe.get("spesa_totale_ricetta"))

    def _num_from_text_local(val: Any) -> Optional[float]:
        if val is None:
            return None
        s = str(val).strip().lower()
        if not s:
            return None
        s = s.replace("€", "").replace("eur", "").replace(",", ".")
        m = re.search(r"-?\d+(?:\.\d+)?", s)
        if not m:
            return None
        try:
            return float(m.group(0))
        except Exception:
            return None

    nt_existing = recipe.get("nutrition_table")
    if not isinstance(nt_existing, dict):
        nt100: Dict[str, Any] = {}
        nttot: Dict[str, Any] = {}
        ntpor: Dict[str, Any] = {}

        def _set_if(block: Dict[str, Any], key: str, *fields: str) -> None:
            for f in fields:
                v = _num_from_text_local(recipe.get(f))
                if v is not None:
                    block[key] = v
                    return

        _set_if(nttot, "energia", "energia_totale", "kcal_ricetta", "kcal_totale", "energia_ricetta")
        _set_if(nt100, "energia", "energia_100g", "kcal_100g")
        _set_if(ntpor, "energia", "energia_porzione", "kcal_porzione", "kcal_per_porzione")

        if nt100 or nttot or ntpor:
            recipe["nutrition_table"] = {"100g": nt100, "totale": nttot, "porzione": ntpor}
            nt_existing = recipe.get("nutrition_table")

    has_nutrition_table = False
    if isinstance(nt_existing, dict):
        for scope in ("100g", "totale", "porzione"):
            block = nt_existing.get(scope)
            if isinstance(block, dict):
                for v in block.values():
                    if v is not None and not _is_zero_like(v):
                        has_nutrition_table = True
                        break
            if has_nutrition_table:
                break

    provider = None
    try:
        ai_comp = recipe.get("ai_completion")
        if isinstance(ai_comp, dict):
            provider = ai_comp.get("provider")
    except Exception:
        provider = None
    use_local_ai = allow_local_ai and not (
        isinstance(provider, str) and provider.strip().lower() not in {"", "ollama", "local"}
    )

    def _nutrition_has_values(nut: Dict[str, Any]) -> bool:
        if not isinstance(nut, dict):
            return False
        for scope in ("total", "per_100g", "per_portion"):
            block = nut.get(scope)
            if not isinstance(block, dict):
                continue
            for v in block.values():
                try:
                    if float(v) > 0:
                        return True
                except Exception:
                    continue
        return False

    def _cost_has_values(cost: Dict[str, Any]) -> bool:
        if not isinstance(cost, dict):
            return False
        total = cost.get("total_cost")
        try:
            if total is not None and float(total) > 0:
                return True
        except Exception:
            pass
        items = cost.get("items")
        if isinstance(items, list):
            for it in items:
                if not isinstance(it, dict):
                    continue
                val = it.get("cost")
                try:
                    if val is not None and float(val) > 0:
                        return True
                except Exception:
                    continue
        return False

    if price_engine is not None and not (has_cost_lines or has_cost_tot):
        try:
            cost_obj = price_engine.compute_cost_safe(recipe)
            cost_summary = _coerce_str_dict(cost_obj)
        except Exception:
            cost_summary = {}
    if cost_summary and not _cost_has_values(cost_summary):
        cost_summary = {}

    if nutrition_engine is not None and not has_nutrition_table:
        try:
            nobj = nutrition_engine.compute_nutrition_safe(recipe)
            nutrition_dict = _coerce_str_dict(nobj)
        except Exception:
            nutrition_dict = {}
    if nutrition_dict and not _nutrition_has_values(nutrition_dict):
        nutrition_dict = {}

    if not cost_summary and PricesDB and not (has_cost_lines or has_cost_tot):
        try:
            pdb = PricesDB.load(base_dir / "data/prices/prezzi_ingredienti_compilati.json")
        except Exception:
            pdb = None
        total_cost = 0.0
        cost_lines = []
        if pdb:
            for ing in recipe.get("ingredients", []):
                try:
                    q_val = ing.get("qty", 0)
                    q = float(q_val if q_val is not None else 0)
                except Exception:
                    q = 0.0
                u = ing.get("unit", "")
                n = ing.get("name", "")
                if q > 0:
                    res = pdb.cost_for_quantity(n, q, u)
                    if res.get("ok"):
                        val = res.get("cost_eur", 0.0)
                        total_cost += float(val)
                        cost_lines.append({"ingrediente": n, "prezzo_calcolato": f"{val:.2f}", "scarto": ""})
        recipe["cost_lines"] = cost_lines
        recipe["spesa_totale_ricetta"] = f"{total_cost:.2f}"
    else:
        if cost_summary:
            recipe["cost_summary"] = cost_summary
            recipe["cost_lines"] = _cost_lines_from_summary(cost_summary)
            if isinstance(cost_summary, dict):
                total_cost = cost_summary.get("total_cost")
                if isinstance(total_cost, (int, float)) and total_cost > 0:
                    recipe["spesa_totale_ricetta"] = f"{total_cost:.2f}"

    cost_lines = recipe.get("cost_lines")
    if isinstance(cost_lines, list):
        def _scarto_empty_or_zero(val: Any) -> bool:
            if _is_empty_value(val):
                return True
            s = str(val).strip().replace("%", "").strip()
            if not s:
                return True
            return _is_zero_like(s) or s == "0"

        for row in cost_lines:
            if not isinstance(row, dict):
                continue
            scarto_raw = row.get("scarto") or row.get("scarto_pct") or row.get("waste_pct")
            if not _scarto_empty_or_zero(scarto_raw):
                continue
            name = row.get("ingrediente") or row.get("ingredient") or row.get("ingredient_raw") or row.get("name")
            est = _estimate_waste_pct(name)
            if est is not None and est != "0":
                row["scarto"] = est

    def _num_from_text(val: Any) -> Optional[float]:
        if val is None:
            return None
        s = str(val).strip().lower()
        if not s:
            return None
        s = s.replace("€", "").replace("eur", "").replace(",", ".")
        m = re.search(r"-?\d+(?:\.\d+)?", s)
        if not m:
            return None
        try:
            return float(m.group(0))
        except Exception:
            return None

    if _is_empty_value(recipe.get("spesa_totale_ricetta")):
        cost_sum = 0.0
        has_vals = False
        for row in recipe.get("cost_lines", []) or []:
            if not isinstance(row, dict):
                continue
            val = _num_from_text(row.get("prezzo_calcolato") or row.get("cost"))
            if val is None:
                continue
            cost_sum += float(val)
            has_vals = True
        if has_vals:
            recipe["spesa_totale_ricetta"] = f"{cost_sum:.2f}"

    if _is_empty_value(recipe.get("spesa_totale_acquisto")):
        acq_sum = 0.0
        has_vals = False
        for row in recipe.get("cost_lines", []) or []:
            if not isinstance(row, dict):
                continue
            val = _num_from_text(row.get("prezzo_alimento_acquisto"))
            if val is None:
                continue
            acq_sum += float(val)
            has_vals = True
        if has_vals:
            recipe["spesa_totale_acquisto"] = f"{acq_sum:.2f}"

    if nutrition_dict and not has_nutrition_table:
        recipe["nutrition"] = nutrition_dict
        recipe["nutrition_table"] = _nutrition_table_from(nutrition_dict)
        total_energy = recipe["nutrition_table"].get("totale", {}).get("energia")
        if total_energy is not None:
            recipe["energia_totale"] = int(round(float(total_energy)))

    total_cost_num = None
    try:
        total_cost_num = float(recipe.get("spesa_totale_ricetta", 0) or 0)
    except Exception:
        total_cost_num = 0.0

    if (
        use_local_ai
        and not nutrition_dict
        and not has_nutrition_table
    ):
        ing_txt = ", ".join(
            [f"{i.get('qty')} {i.get('unit')} {i.get('name')}" for i in recipe.get("ingredients", [])]
        )
        ai_data = estimate_nutrition_ai(ing_txt)
        if ai_data:
            if total_cost_num is None or total_cost_num < 0.1:
                total_cost_num = float(ai_data.get("cost_eur", 0) or 0)
                recipe["spesa_totale_ricetta"] = f"{total_cost_num:.2f}"
            kcal_val = float(ai_data.get("kcal", 0) or 0)
            recipe["energia_totale"] = int(kcal_val or 0)
            nutrition_dict = {
                "total": {
                    "kcal": kcal_val,
                    "carbs_g": ai_data.get("carbs", 0) or 0,
                    "fat_g": ai_data.get("fats", 0) or 0,
                    "protein_g": ai_data.get("proteins", 0) or 0,
                }
            }
            recipe["nutrition"] = nutrition_dict
            recipe["nutrition_table"] = _nutrition_table_from(nutrition_dict)

    if not has_nutrition_table:
        def _estimate_kcal_fallback() -> Dict[str, Any]:
            total_weight_g = 0.0
            for ing in recipe.get("ingredients", []) or []:
                if not isinstance(ing, dict):
                    continue
                qty = _num_from_text_local(ing.get("qty"))
                unit = str(ing.get("unit") or "").strip().lower()
                if qty is None or qty <= 0:
                    continue
                if unit in {"g", "gr", "grammi", "grammo"}:
                    total_weight_g += qty
                elif unit in {"kg", "kilo", "kilogrammi", "kilogrammo"}:
                    total_weight_g += qty * 1000.0
                elif unit in {"mg"}:
                    total_weight_g += qty / 1000.0
                elif unit in {"ml"}:
                    total_weight_g += qty
                elif unit in {"l", "lt", "litro", "litri"}:
                    total_weight_g += qty * 1000.0
                elif unit in {"cl"}:
                    total_weight_g += qty * 10.0
                elif unit in {"dl"}:
                    total_weight_g += qty * 100.0

            base_kcal_100g = 150.0
            if total_weight_g > 0:
                kcal_tot = (total_weight_g / 100.0) * base_kcal_100g
                kcal_100 = base_kcal_100g
            else:
                kcal_tot = 400.0
                kcal_100 = None

            servings = _num_from_text_local(recipe.get("servings") or recipe.get("porzioni"))
            kcal_por = None
            if servings and servings > 0:
                kcal_por = kcal_tot / servings

            nt = {"100g": {}, "totale": {}, "porzione": {}}
            nt["totale"]["energia"] = kcal_tot
            if kcal_100 is not None:
                nt["100g"]["energia"] = kcal_100
            if kcal_por is not None:
                nt["porzione"]["energia"] = kcal_por
            return {"kcal_tot": kcal_tot, "kcal_100": kcal_100, "kcal_por": kcal_por, "table": nt}

        est = _estimate_kcal_fallback()
        recipe["nutrition_table"] = est["table"]
        recipe["energia_totale"] = int(round(float(est["kcal_tot"])) if est["kcal_tot"] is not None else 0)
        if est.get("kcal_por") is not None:
            recipe["energia_porzione"] = float(est["kcal_por"])

    if allergens is not None and _is_empty_value(recipe.get("allergens")) and _is_empty_value(recipe.get("allergens_text")):
        try:
            summary = allergens.infer_allergens(recipe)
            summary_dict = summary.to_dict() if hasattr(summary, "to_dict") else {}
        except Exception:
            summary_dict = {}
        if summary_dict:
            present = summary_dict.get("present") or []
            if not isinstance(present, list):
                present = [present]
            present_codes = [str(x).strip() for x in present if str(x).strip()]

            traces = summary_dict.get("traces") or []
            if not isinstance(traces, list):
                traces = [traces]
            trace_codes = [str(x).strip() for x in traces if str(x).strip()]

            recipe["allergens"] = present_codes
            recipe["allergens_text"] = ", ".join(present_codes)
            recipe["traces_allergens"] = ", ".join(trace_codes)

            df_raw = summary_dict.get("diet_flags")
            df = df_raw if isinstance(df_raw, dict) else {}
            recipe["diet_flags"] = {
                "vegetarian": bool(df.get("vegetarian", {}).get("value", False)),
                "vegan": bool(df.get("vegan", {}).get("value", False)),
                "gluten_free": bool(df.get("gluten_free", {}).get("value", False)),
                "lactose_free": bool(df.get("lactose_free", {}).get("value", False)),
            }

    if equipment is not None and _is_empty_value(recipe.get("equipment_text")) and _is_empty_value(recipe.get("equipment")):
        try:
            root = base_dir / "data" / "equipment_pasticceria.json"
            if root.exists():
                eq_db = equipment.EquipmentDB.load(root)
                eq = eq_db.suggest_from_text(source_text, max_items=20)
                if isinstance(eq, dict) and eq.get("ok"):
                    raw_names = [x.get("name") for x in eq.get("suggested", []) if isinstance(x, dict)]
                    names = [str(name).strip() for name in raw_names if name]
                    recipe["equipment"] = names
                    recipe["equipment_text"] = ", ".join(names)
        except Exception:
            pass

    try:
        servings = float(recipe.get("servings", 0) or 0)
    except Exception:
        servings = 0.0
    try:
        total_cost_num = float(recipe.get("spesa_totale_ricetta", 0) or 0)
    except Exception:
        total_cost_num = 0.0
    if servings > 0:
        recipe["spesa_per_porzione"] = f"{(total_cost_num / servings):.2f}"

    recipe["ingredients_text"] = _build_ingredients_text(recipe.get("ingredients", []))
    recipe["steps_text"] = _build_steps_text(recipe.get("steps", []), numbered=True)
    recipe["steps_text_plain"] = _build_steps_text(recipe.get("steps", []), numbered=False)


# ------------------------------------------------------------
# Template context
# ------------------------------------------------------------
def _normalize_allergen_token(token: str) -> str:
    if allergens is not None and hasattr(allergens, "normalize_text"):
        try:
            return allergens.normalize_text(token)
        except Exception:
            pass
    return re.sub(r"\s+", " ", re.sub(r"[^\w\s/-]+", " ", str(token).lower())).strip()


def _allergen_code_map() -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    labels = getattr(allergens, "ALLERGEN_LABELS_IT", {}) if allergens is not None else {}
    for code, label in (labels or {}).items():
        mapping[_normalize_allergen_token(code)] = code
        mapping[_normalize_allergen_token(label)] = code
    synonyms = {
        "frutta secca": "frutta_a_guscio",
        "frutta a guscio": "frutta_a_guscio",
        "frutta a guscio e derivati": "frutta_a_guscio",
        "frutta secca e derivati": "frutta_a_guscio",
        "noci": "frutta_a_guscio",
        "nocciole": "frutta_a_guscio",
        "mandorle": "frutta_a_guscio",
        "uova e derivati": "uova",
        "latte e derivati": "latte",
        "arachidi e derivati": "arachidi",
        "soia e derivati": "soia",
        "pesce e derivati": "pesce",
        "pesci": "pesce",
        "crostacei e derivati": "crostacei",
        "molluschi e derivati": "molluschi",
        "frutti di mare": "molluschi",
        "lupini e derivati": "lupini",
        "sedano e derivati": "sedano",
        "senape e derivati": "senape",
        "sesamo e derivati": "sesamo",
        "cereali contenenti glutine": "glutine",
        "cereali con glutine": "glutine",
        "grano": "glutine",
        "anidride solforosa": "solfiti",
        "anidride solforosa e solfiti": "solfiti",
        "solforosa": "solfiti",
        "gluten": "glutine",
    }
    for k, v in synonyms.items():
        mapping[_normalize_allergen_token(k)] = v
    return mapping


def _extract_allergen_codes(text: str) -> List[str]:
    if not text:
        return []
    mapping = _allergen_code_map()
    codes: List[str] = []
    for part in re.split(r"[;,/]+", str(text)):
        token = _normalize_allergen_token(part)
        if not token:
            continue
        for key, code in mapping.items():
            if token == key or key in token:
                if code not in codes:
                    codes.append(code)
    return codes


def _build_allergen_icons(recipe: Dict[str, Any]) -> Tuple[List[Dict[str, str]], List[Dict[str, str]]]:
    labels = getattr(allergens, "ALLERGEN_LABELS_IT", {}) if allergens is not None else {}
    present = recipe.get("allergens") or recipe.get("allergens_present") or []
    traces = recipe.get("traces_allergens") or recipe.get("allergens_traces") or []

    if not isinstance(present, list) or not present:
        present = _extract_allergen_codes(recipe.get("allergens_text") or recipe.get("allergeni") or "")
    if not isinstance(traces, list) or not traces:
        traces = _extract_allergen_codes(recipe.get("traces_allergens") or recipe.get("tracce_allergeni") or "")

    def build_list(items: List[Any]) -> List[Dict[str, str]]:
        out: List[Dict[str, str]] = []
        for code in items:
            c = str(code).strip()
            if not c:
                continue
            label = labels.get(c, c)
            out.append({"code": c, "label": label})
        return out

    return build_list(present), build_list(traces)


def clean_recipe_data(recipe: Dict[str, Any]) -> Dict[str, Any]:
    """Pulisce i dati della ricetta per il rendering: trim stringhe, rimozione None, filtro liste.
    
    Operazioni:
    - Converte None in stringa vuota per stringhe
    - Trim delle stringhe
    - Filtra liste rimuovendo None e stringhe vuote
    - Ricorsivo su dizionari annidati
    - Rimuove chiavi con valore vuoto (None, "", [], {})
    
    Args:
        recipe: dizionario della ricetta
    
    Returns:
        Dizionario pulito (nessun None, stringhe trimmate, liste filtrate)
    """
    def _is_empty(val: Any) -> bool:
        """Controlla se un valore è vuoto."""
        if val is None:
            return True
        if isinstance(val, str) and not val.strip():
            return True
        if isinstance(val, (list, dict)) and len(val) == 0:
            return True
        return False

    def _clean(val: Any) -> Any:
        if val is None:
            return None
        if isinstance(val, dict):
            cleaned: Dict[str, Any] = {}
            for k, v in val.items():
                c = _clean(v)
                # Mantieni il valore solo se non è vuoto/None
                if not _is_empty(c):
                    cleaned[k] = c
            return cleaned if cleaned else None
        if isinstance(val, (list, tuple)):
            items: List[Any] = []
            for item in val:
                c = _clean(item)
                if not _is_empty(c):
                    items.append(c)
            return items if items else None
        if isinstance(val, str):
            s = val.strip()
            return s if s else None
        # Mantieni numeri, booleani, date, ecc. come sono
        return val

    if not isinstance(recipe, dict):
        return {}
    
    result = _clean(recipe)
    return result if isinstance(result, dict) else {}


def build_template_context(recipe: Dict[str, Any]) -> Dict[str, Any]:
    ctx: Dict[str, Any] = dict(recipe or {})
    ctx["recipe"] = dict(recipe or {})

    def _clean_text_value(val: Any) -> str:
        if val is None:
            return ""
        s = str(val).strip()
        if not s:
            return ""
        if s.lower() in {"none", "null", "n/d", "n.d", "nd", "n.a", "n/a"}:
            return ""
        return s

    def _clean_title_text(val: Any) -> str:
        s = _clean_text_value(val)
        if not s:
            return ""
        s = re.sub(r"\bTitolo\s*:\s*", "", s, flags=re.IGNORECASE).strip()
        if "|" in s:
            parts = [p.strip() for p in s.split("|") if p.strip()]
            if parts:
                uniq: List[str] = []
                for p in parts:
                    if p.lower() not in [u.lower() for u in uniq]:
                        uniq.append(p)
                s = uniq[0] if uniq else s
        return s

    def _clean_num_value(val: Any) -> Any:
        if val is None:
            return ""
        if isinstance(val, str):
            s = val.strip()
            if not s:
                return ""
            if s.lower() in {"none", "null", "n/d", "n.d", "nd", "n.a", "n/a"}:
                return ""
            return s
        return val

    def _to_float(val: Any) -> Optional[float]:
        if val is None:
            return None
        if isinstance(val, (int, float)):
            return float(val)
        if isinstance(val, str):
            s = val.strip()
            if not s:
                return None
            s = s.replace("€", "")
            s = re.sub(r"(?i)eur", "", s)
            s = s.replace(" ", "")
            s = s.replace(",", ".")
            try:
                return float(s)
            except Exception:
                return None
        return None

    def _normalize_unit_text(val: Any) -> str:
        s = _clean_text_value(val)
        if not s:
            return ""
        return re.sub(r"\b(ud|u|unita|unità|unit)\b", "pz", s, flags=re.IGNORECASE)

    def _format_temp(val: Any) -> str:
        s = _clean_text_value(val)
        if not s:
            return ""
        if "°" in s or re.search(r"\b[cf]\b", s.lower()):
            return s
        return f"{s} °C"

    def _format_currency(val: Any) -> str:
        s = _clean_text_value(val)
        if not s:
            return ""
        if "€" in s or "eur" in s.lower():
            return s
        if re.search(r"\d", s):
            return f"{s} €"
        return s

    def _format_price_unit(val: Any) -> str:
        s = _clean_text_value(val)
        if not s:
            return ""
        s = s.replace("?/","€/")
        if "/" in s and "€" not in s and "eur" not in s.lower():
            s = s.replace("/", "€/", 1)
        if "€" not in s and "eur" not in s.lower() and re.search(r"\d", s):
            s = f"{s} €"
        return s

    def _listify(val: Any) -> str:
        if val is None:
            return ""
        items: List[str] = []
        if isinstance(val, list):
            items = [str(x).strip() for x in val if str(x).strip()]
        else:
            s = _clean_text_value(val)
            if not s:
                return ""
            if "\n" in s:
                items = [ln.strip() for ln in s.splitlines() if ln.strip()]
            else:
                items = [p.strip() for p in re.split(r"[;,]+", s) if p.strip()]
        cleaned: List[str] = []
        for item in items:
            item = re.sub(r"^\d+[\.\)]\s*", "", item)
            item = item.strip("-•* \t").strip()
            item = re.sub(
                r"^(?:attrezzature\s*)?(semplici|generiche|professionali|specifiche|pasticceria)\s*:\s*",
                "",
                item,
                flags=re.IGNORECASE,
            )
            if item:
                cleaned.append(f"- {item}")
        return "\n".join(cleaned)

    def _norm_text(val: Any) -> str:
        s = _clean_text_value(val).lower()
        s = re.sub(r"[^\w\s-]+", " ", s)
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def _diet_has(text: str, keys: List[str]) -> bool:
        for kw in keys:
            if " " in kw:
                if kw in text:
                    return True
            else:
                if re.search(r"\b" + re.escape(kw), text):
                    return True
        return False

    title = _clean_title_text(recipe.get("title") or recipe.get("titolo")) or "Ricetta"
    category = _clean_text_value(recipe.get("category") or recipe.get("categoria"))
    subcategory = _clean_text_value(recipe.get("subcategory") or recipe.get("sottocategoria"))
    servings = _clean_text_value(recipe.get("servings") or recipe.get("porzioni"))
    difficulty = _clean_text_value(recipe.get("difficulty") or recipe.get("difficolta"))

    peso_totale_ricetta_g = _to_float(recipe.get("peso_totale_ricetta_g") or recipe.get("peso_totale_ricetta"))
    resa_totale = _clean_text_value(recipe.get("resa_totale") or recipe.get("yield") or recipe.get("resa"))

    prep = recipe.get("prep_time_min")
    cook = recipe.get("cook_time_min")
    total = recipe.get("total_time_min")

    def _time_str(v: Any) -> str:
        try:
            vi = int(float(v))
            return f"{vi} min"
        except Exception:
            return ""

    def _parse_minutes(raw: str) -> Optional[int]:
        s = (raw or "").strip().lower()
        if not s:
            return None
        h = 0
        m = 0
        mh = re.search(r"(\d+)\s*h", s)
        if mh:
            try:
                h = int(mh.group(1))
            except Exception:
                h = 0
        mm = re.search(r"(\d+)\s*(min|m)\b", s)
        if mm:
            try:
                m = int(mm.group(1))
            except Exception:
                m = 0
        if h or m:
            return h * 60 + m
        mn = re.search(r"(\d+)", s)
        if mn:
            try:
                return int(mn.group(1))
            except Exception:
                return None
        return None

    def _extract_time(label: str, raw: str) -> Optional[int]:
        m = re.search(rf"{label}\s*[:\-]?\s*([0-9][^,;]*)", (raw or "").lower())
        if not m:
            return None
        return _parse_minutes(m.group(1))

    prep_s = _time_str(prep)
    if not prep_s:
        prep_s = _clean_text_value(recipe.get("tempo_preparazione") or "")
    cook_s = _time_str(cook)
    if not cook_s:
        cook_s = _clean_text_value(recipe.get("tempo_cottura") or "")
    total_s = _time_str(total)
    if not total_s:
        total_s = _clean_text_value(recipe.get("tempo_totale") or "")

    tempo_raw = _clean_text_value(recipe.get("tempo_dettaglio") or recipe.get("tempo"))
    if tempo_raw and (not prep_s or not cook_s or not total_s):
        p2 = _extract_time(r"(?:prep|preparazione)", tempo_raw)
        c2 = _extract_time(r"(?:cottura|cook)", tempo_raw)
        t2 = _extract_time(r"(?:totale|complessivo|total)", tempo_raw)
        if not prep_s and p2 is not None:
            prep_s = f"{p2} min"
        if not cook_s and c2 is not None:
            cook_s = f"{c2} min"
        if not total_s and t2 is not None:
            total_s = f"{t2} min"

    tempo_parts = []
    if prep_s:
        tempo_parts.append(f"Prep {prep_s}")
    if cook_s:
        tempo_parts.append(f"Cottura {cook_s}")
    if total_s:
        tempo_parts.append(f"Tot {total_s}")
    tempo_dettaglio = ", ".join([p for p in tempo_parts if p.strip()])
    if not tempo_dettaglio and tempo_raw:
        tempo_dettaglio = tempo_raw

    ingredients_text = _build_ingredients_text(recipe.get("ingredients", []))
    if not ingredients_text:
        ingredients_text = recipe.get("ingredients_text") or ""
    steps_text = _build_steps_text(recipe.get("steps", []), numbered=True)
    steps_text_plain = _build_steps_text(recipe.get("steps", []), numbered=False)
    if not steps_text:
        steps_text = recipe.get("steps_text") or ""
    if not steps_text_plain:
        steps_text_plain = recipe.get("steps_text_plain") or recipe.get("steps_text") or ""

    vegetariano_raw = _clean_text_value(recipe.get("vegetariano_flag") or recipe.get("vegetariano flag") or "")

    diet_flags_raw = recipe.get("diet_flags")
    diet_flags = diet_flags_raw if isinstance(diet_flags_raw, dict) else {}
    diet_labels = []
    if diet_flags.get("vegan"):
        diet_labels.append("vegana")
    elif diet_flags.get("vegetarian"):
        diet_labels.append("vegetariana")
    if diet_flags.get("gluten_free"):
        diet_labels.append("senza glutine")
    if diet_flags.get("lactose_free"):
        diet_labels.append("senza lattosio")

    diet_text_raw = _clean_text_value(
        recipe.get("diet_text")
        or recipe.get("diete_text")
        or recipe.get("diete")
        or recipe.get("diet")
        or recipe.get("dieta")
        or ""
    )
    diet_text_norm = _norm_text(" ".join([diet_text_raw, vegetariano_raw, " ".join(diet_labels)]))

    vegetariano_flag = vegetariano_raw
    if not vegetariano_flag:
        if diet_flags.get("vegan") or diet_flags.get("vegetarian"):
            vegetariano_flag = "Si"
        elif diet_flags:
            vegetariano_flag = "No"

    vegan_ok = bool(diet_flags.get("vegan")) or _diet_has(diet_text_norm, ["vegana", "vegano", "vegan"])
    vegetarian_ok = bool(diet_flags.get("vegetarian")) or _diet_has(diet_text_norm, ["vegetar"])
    if vegan_ok:
        vegetarian_ok = True

    vegetariano_ok = False
    vegetariano_norm = _norm_text(vegetariano_flag)
    if vegetariano_norm in {"no", "false", "0"}:
        vegetariano_ok = False
    elif vegetariano_norm in {"si", "sì", "yes", "true", "1"}:
        vegetariano_ok = True
    elif vegetariano_norm:
        vegetariano_ok = vegetarian_ok or vegan_ok
    else:
        vegetariano_ok = vegetarian_ok or vegan_ok

    diete_scelta: List[str] = []
    if vegetarian_ok:
        diete_scelta.append("Dieta vegetariana")
    if vegan_ok:
        diete_scelta.append("Dieta vegana")
    if _diet_has(diet_text_norm, ["plant-based", "plant based", "plantbased", "base vegetale", "a base vegetale"]):
        diete_scelta.append("Dieta plant-based")
    if _diet_has(diet_text_norm, ["flexitar", "flexitarian"]):
        diete_scelta.append("Dieta flexitariana")
    if _diet_has(diet_text_norm, ["pescetar", "pescetarian", "pescatar"]):
        diete_scelta.append("Dieta pescetariana")

    diete_cliniche: List[str] = []
    if diet_flags.get("gluten_free") or _diet_has(
        diet_text_norm, ["senza glutine", "gluten free", "gluten-free", "celiach", "celiac"]
    ):
        diete_cliniche.append("Dieta senza glutine (per celiachia)")
    if _diet_has(diet_text_norm, ["ipocalor", "low calorie", "basso apporto calorico"]):
        diete_cliniche.append("Dieta ipocalorica")
    if _diet_has(diet_text_norm, ["ipercalor", "high calorie", "alto apporto calorico"]):
        diete_cliniche.append("Dieta ipercalorica")
    if _diet_has(diet_text_norm, ["diabet"]):
        diete_cliniche.append("Dieta per diabetici")
    if _diet_has(diet_text_norm, ["ipertens", "basso sodio", "low sodium"]):
        diete_cliniche.append("Dieta per ipertensione")
    if _diet_has(diet_text_norm, ["reflusso", "gastroesofageo", "reflux"]):
        diete_cliniche.append("Dieta per reflusso gastroesofageo")

    diete_culturali: List[str] = []
    if _diet_has(diet_text_norm, ["halal", "hallal"]):
        diete_culturali.append("Dieta halal")
    if _diet_has(diet_text_norm, ["kosher", "kasher"]):
        diete_culturali.append("Dieta kosher")
    if _diet_has(diet_text_norm, ["induista", "hindu"]):
        diete_culturali.append("Dieta induista")
    if _diet_has(diet_text_norm, ["buddh", "buddis"]):
        diete_culturali.append("Dieta buddhista")

    diete_stile: List[str] = []
    if _diet_has(diet_text_norm, ["mediterran"]):
        diete_stile.append("Dieta mediterranea")
    if _diet_has(diet_text_norm, ["chetogen", "keto"]):
        diete_stile.append("Dieta chetogenica")

    allergens_text = _clean_text_value(recipe.get("allergens_text"))
    allergens_list = recipe.get("allergens")
    if not allergens_text and isinstance(allergens_list, list):
        labels = getattr(allergens, "ALLERGEN_LABELS_IT", {}) if allergens is not None else {}
        allergens_text = ", ".join([labels.get(str(x), str(x)) for x in allergens_list if x])

    traces_raw = recipe.get("traces_allergens") or recipe.get("tracce_allergeni") or ""
    if isinstance(traces_raw, list):
        labels = getattr(allergens, "ALLERGEN_LABELS_IT", {}) if allergens is not None else {}
        traces_text = ", ".join([labels.get(str(x), str(x)) for x in traces_raw if x])
    else:
        traces_text = _clean_text_value(traces_raw)

    # Fallback per allergeni e tracce se non presenti
    if not allergens_text:
        allergens_text = "Nessun allergene"
    if not traces_text:
        traces_text = "Nessuna traccia"
    equipment_text = _clean_text_value(recipe.get("equipment_text"))
    equipment_list = recipe.get("equipment")
    if not equipment_text and isinstance(equipment_list, list):
        equipment_text = ", ".join([str(x) for x in equipment_list if x])

    vino_descrizione = _clean_text_value(recipe.get("vino_descrizione") or recipe.get("wine_pairing") or "")
    vino_temp = _format_temp(
        recipe.get("vino_temperatura_servizio") or recipe.get("vino temperatura servizio") or ""
    )
    vino_regione = _clean_text_value(recipe.get("vino_regione") or recipe.get("vino regione") or "")
    vino_annata = _clean_text_value(recipe.get("vino_annata") or recipe.get("vino annata") or "")
    vino_motivo_annata = _clean_text_value(
        recipe.get("vino_motivo_annata") or recipe.get("vino motivo annata") or ""
    )
    conservazione = _clean_text_value(recipe.get("conservazione") or recipe.get("storage") or "")
    presentazione = _clean_text_value(recipe.get("presentazione_impiattamento") or recipe.get("presentazione") or "")
    attrezzature_semplici = _clean_text_value(
        recipe.get("attrezzature_semplici")
        or recipe.get("attrezzature semplici")
        or recipe.get("equipment_simple")
        or ""
    )
    attrezzature_professionali = _clean_text_value(
        recipe.get("attrezzature_professionali")
        or recipe.get("attrezzature professionali")
        or recipe.get("equipment_professional")
        or ""
    )
    attrezzature_pasticceria = _clean_text_value(
        recipe.get("attrezzature_pasticceria")
        or recipe.get("attrezzature pasticceria")
        or recipe.get("equipment_pasticceria")
        or ""
    )
    attrezzature_generiche = _clean_text_value(
        recipe.get("attrezzature_generiche") or attrezzature_semplici or equipment_text
    )
    attrezzature_specifiche = _clean_text_value(
        recipe.get("attrezzature_specifiche") or attrezzature_professionali or ""
    )
    attrezzature_generiche = _listify(attrezzature_generiche)
    attrezzature_specifiche = _listify(attrezzature_specifiche)
    attrezzature_pasticceria = _listify(attrezzature_pasticceria)

    componenti_raw = recipe.get("ingredienti_componenti")
    componenti_list = componenti_raw if isinstance(componenti_raw, list) else []
    ingredienti_componenti: List[Dict[str, Any]] = []
    for group in componenti_list:
        if not isinstance(group, dict):
            continue
        g_name = _clean_text_value(group.get("name") or group.get("nome"))
        items_raw = group.get("items") if isinstance(group.get("items"), list) else []
        items_clean: List[Dict[str, Any]] = []
        for item in (items_raw or []):
            if not isinstance(item, dict):
                continue
            itm: Dict[str, Any] = {}
            itm["nome"] = _clean_text_value(
                item.get("nome") or item.get("name") or item.get("parsed_name") or item.get("line")
            )
            itm["name"] = itm.get("nome") or ""
            itm["parsed_name"] = _clean_text_value(item.get("parsed_name"))
            itm["line"] = _clean_text_value(item.get("line"))
            itm["quantita"] = item.get("quantita") if item.get("quantita") not in (None, "") else item.get("grams")
            itm["quantita_raw"] = _clean_text_value(item.get("quantita_raw") or item.get("quantita_usata"))
            itm["grams"] = item.get("grams")
            itm["unita"] = _clean_text_value(item.get("unita") or item.get("unit"))
            itm["unit"] = itm.get("unita") or ""
            itm["note"] = _clean_text_value(item.get("note"))
            itm["costo_ingrediente"] = _to_float(item.get("costo_ingrediente"))
            itm["costo_per_qty"] = _to_float(item.get("costo_per_qty"))
            itm["costo_usato"] = _to_float(item.get("costo_usato"))
            itm["costo_spesa"] = _to_float(item.get("costo_spesa"))
            itm["costo_acquisto_minimo"] = _to_float(item.get("costo_acquisto_minimo"))
            itm["acquisto_minimo_g"] = _to_float(item.get("acquisto_minimo_g"))
            items_clean.append(itm)
        ingredienti_componenti.append({"name": g_name, "items": items_clean})

    stagionalita = _clean_text_value(
        recipe.get("stagionalita") or recipe.get("stagionalità") or recipe.get("seasonality") or ""
    )

    allergeni_loghi, allergeni_tracce_loghi = _build_allergen_icons(recipe)

    diete_scelta_text = _listify(diete_scelta)
    diete_cliniche_text = _listify(diete_cliniche)
    diete_culturali_text = _listify(diete_culturali)
    diete_stile_text = _listify(diete_stile)

    ctx.update(
        {
            "titolo": title,
            "title": title,
            "categoria": category,
            "category": category,
            "sottocategoria": subcategory,
            "subcategory": subcategory,
            "porzioni": servings,
            "porzione_standard_display": servings or "n/d",
            "difficolta": difficulty,
            "tempo_preparazione": prep_s,
            "tempo_cottura": cook_s,
            "tempo_totale": total_s,
            "tempo_dettaglio": tempo_dettaglio,
            "tempo": tempo_dettaglio,
            "tempo_riposo": _clean_text_value(recipe.get("tempo_riposo")),
            "resa_totale": resa_totale,
            "peso_totale_ricetta_g": peso_totale_ricetta_g or 0,
            "codice_ricetta": _clean_text_value(recipe.get("codice_ricetta")),
            "metodo_principale": _clean_text_value(recipe.get("metodo_principale")),
            "image_src": _clean_text_value(recipe.get("image_src") or recipe.get("image")),
            "ingredienti": ingredients_text,
            "ingredienti_blocco": ingredients_text,
            "ingredienti_varianti": _clean_text_value(recipe.get("ingredienti_varianti")),
            "ingredienti_componenti": ingredienti_componenti,
            "procedimento": steps_text,
            "procedimento_blocco": steps_text,
            "procedimento_blocco_plain": steps_text_plain,
            "procedimento_parametri_tecnici": _clean_text_value(recipe.get("procedimento_parametri_tecnici")),
            "procedimento_punti_critici": _clean_text_value(recipe.get("procedimento_punti_critici")),
            "procedimento_errori_tipici": _clean_text_value(recipe.get("procedimento_errori_tipici")),
            "allergeni_elenco": allergens_text,
            "allergeni_text": allergens_text,
            "tracce_allergeni": traces_text,
            "vino_descrizione": vino_descrizione or "",
            "vino_temperatura_servizio": vino_temp or "",
            "vino_regione": vino_regione or "",
            "vino_annata": vino_annata or "",
            "vino_motivo_annata": vino_motivo_annata or "",
            "conservazione": conservazione or "",
            "note_haccp": _clean_text_value(recipe.get("note_haccp")) or "",
            "presentazione_impiattamento": presentazione or "",
            "attrezzature": equipment_text or "",
            "attrezzature_text": equipment_text or "",
            "attrezzature_generiche": attrezzature_generiche or "",
            "attrezzature_specifiche": attrezzature_specifiche or "",
            "attrezzature_semplici": attrezzature_semplici or "",
            "attrezzature_professionali": attrezzature_professionali or "",
            "attrezzature_pasticceria": attrezzature_pasticceria or "",
            "vegetariano_flag": vegetariano_flag,
            "vegetariano_ok": vegetariano_ok,
            "diete_text": ", ".join(diet_labels) if diet_labels else "",
            "diete_scelta_alimentare": diete_scelta_text,
            "diete_cliniche": diete_cliniche_text,
            "diete_culturali": diete_culturali_text,
            "diete_stile": diete_stile_text,
            "diete_note": _clean_text_value(recipe.get("diete_note")),
            "stagionalita": stagionalita or "",
            "allergeni_loghi": allergeni_loghi,
            "allergeni_tracce_loghi": allergeni_tracce_loghi,
            "abbattimento_raffreddamento": _clean_text_value(recipe.get("abbattimento_raffreddamento")),
            "rigenerazione_servizio": _clean_text_value(recipe.get("rigenerazione_servizio")),
            "porzionatura_servizio": _clean_text_value(recipe.get("porzionatura_servizio")),
            "abbinamenti": _clean_text_value(recipe.get("abbinamenti")),
            "decorazioni_garnish": _clean_text_value(recipe.get("decorazioni_garnish")),
            "autore": _clean_text_value(recipe.get("autore")),
            "data_creazione": _clean_text_value(recipe.get("data_creazione")),
            "data_modifica": _clean_text_value(recipe.get("data_modifica")),
            "tag_parole_chiave": _clean_text_value(recipe.get("tag_parole_chiave")),
            "nome_file_standardizzato": _clean_text_value(recipe.get("nome_file_standardizzato")),
            "note_interne": _clean_text_value(recipe.get("note_interne")),
        }
    )

    nt = recipe.get("nutrition_table") if isinstance(recipe.get("nutrition_table"), dict) else {}
    nt100 = nt.get("100g", {}) if isinstance(nt, dict) else {}
    nttot = nt.get("totale", {}) if isinstance(nt, dict) else {}
    ntpor = nt.get("porzione", {}) if isinstance(nt, dict) else {}

    def _num(v: Any) -> float:
        try:
            return float(v)
        except Exception:
            return 0.0

    ctx.update(
        {
            "energia_100g": _clean_num_value(nt100.get("energia")),
            "energia_totale": _clean_num_value(nttot.get("energia")),
            "energia_porzione": _clean_num_value(ntpor.get("energia")),
            "energia_ricetta": _clean_num_value(nttot.get("energia")),
            "carboidrati_totali_100g": _clean_num_value(nt100.get("carboidrati_totali")),
            "carboidrati_totali_totale": _clean_num_value(nttot.get("carboidrati_totali")),
            "carboidrati_100g": _clean_num_value(nt100.get("carboidrati_totali")),
            "carboidrati_ricetta": _clean_num_value(nttot.get("carboidrati_totali")),
            "carboidrati_porzione": _clean_num_value(ntpor.get("carboidrati_totali")),
            "di_cui_zuccheri_100g": _clean_num_value(nt100.get("di_cui_zuccheri")),
            "di_cui_zuccheri_totale": _clean_num_value(nttot.get("di_cui_zuccheri")),
            "zuccheri_100g": _clean_num_value(nt100.get("di_cui_zuccheri")),
            "zuccheri_ricetta": _clean_num_value(nttot.get("di_cui_zuccheri")),
            "zuccheri_porzione": _clean_num_value(ntpor.get("di_cui_zuccheri")),
            "grassi_totali_100g": _clean_num_value(nt100.get("grassi_totali")),
            "grassi_totali_totale": _clean_num_value(nttot.get("grassi_totali")),
            "grassi_totali_porzione": _clean_num_value(ntpor.get("grassi_totali")),
            "grassi_totali_ricetta": _clean_num_value(nttot.get("grassi_totali")),
            "di_cui_saturi_100g": _clean_num_value(nt100.get("di_cui_saturi")),
            "di_cui_saturi_totale": _clean_num_value(nttot.get("di_cui_saturi")),
            "di_cui_grassi_saturi_100g": _clean_num_value(nt100.get("di_cui_saturi")),
            "di_cui_grassi_saturi_totale": _clean_num_value(nttot.get("di_cui_saturi")),
            "grassi_saturi_100g": _clean_num_value(nt100.get("di_cui_saturi")),
            "grassi_saturi_ricetta": _clean_num_value(nttot.get("di_cui_saturi")),
            "grassi_saturi_porzione": _clean_num_value(ntpor.get("di_cui_saturi")),
            "proteine_totali_100g": _clean_num_value(nt100.get("proteine_totali")),
            "proteine_totali_totale": _clean_num_value(nttot.get("proteine_totali")),
            "proteine_100g": _clean_num_value(nt100.get("proteine_totali")),
            "proteine_totale": _clean_num_value(nttot.get("proteine_totali")),
            "proteine_ricetta": _clean_num_value(nttot.get("proteine_totali")),
            "proteine_porzione": _clean_num_value(ntpor.get("proteine_totali")),
            "fibre_100g": _clean_num_value(nt100.get("fibre")),
            "fibre_totale": _clean_num_value(nttot.get("fibre")),
            "fibre_ricetta": _clean_num_value(nttot.get("fibre")),
            "fibre_porzione": _clean_num_value(ntpor.get("fibre")),
            "fibra_alimentare_100g": _clean_num_value(nt100.get("fibre")),
            "fibra_alimentare_totale": _clean_num_value(nttot.get("fibre")),
            "sodio_100g": _clean_num_value(nt100.get("sodio")),
            "sodio_totale": _clean_num_value(nttot.get("sodio")),
            "sodio_ricetta": _clean_num_value(nttot.get("sodio")),
            "sodio_porzione": _clean_num_value(ntpor.get("sodio")),
            "sale_100g": _clean_num_value(nt100.get("sodio")),
            "sale_totale": _clean_num_value(nttot.get("sodio")),
            "sodio_100g_mg": _clean_num_value(nt100.get("sodio")),
            "sodio_totale_mg": _clean_num_value(nttot.get("sodio")),
            "sodio_ricetta_mg": _clean_num_value(nttot.get("sodio")),
            "sodio_porzione_mg": _clean_num_value(ntpor.get("sodio")),
            "colesterolo_100g": _clean_num_value(nt100.get("colesterolo")),
            "colesterolo_totale": _clean_num_value(nttot.get("colesterolo")),
            "colesterolo_ricetta": _clean_num_value(nttot.get("colesterolo")),
            "colesterolo_porzione": _clean_num_value(ntpor.get("colesterolo")),
            "colesterolo_100g_mg": _clean_num_value(nt100.get("colesterolo")),
            "colesterolo_ricetta_mg": _clean_num_value(nttot.get("colesterolo")),
            "colesterolo_porzione_mg": _clean_num_value(ntpor.get("colesterolo")),
            "grassi_monoinsaturi_100g": _clean_num_value(nt100.get("monoinsaturi")),
            "grassi_monoinsaturi_totale": _clean_num_value(nttot.get("monoinsaturi")),
            "grassi_monoinsaturi_ricetta": _clean_num_value(nttot.get("monoinsaturi")),
            "grassi_monoinsaturi_porzione": _clean_num_value(ntpor.get("monoinsaturi")),
            "grassi_polinsaturi_100g": _clean_num_value(nt100.get("polinsaturi")),
            "grassi_polinsaturi_totale": _clean_num_value(nttot.get("polinsaturi")),
            "grassi_polinsaturi_ricetta": _clean_num_value(nttot.get("polinsaturi")),
            "grassi_polinsaturi_porzione": _clean_num_value(ntpor.get("polinsaturi")),
            "monoinsaturi_100g": _clean_num_value(nt100.get("monoinsaturi")),
            "monoinsaturi_totale": _clean_num_value(nttot.get("monoinsaturi")),
            "polinsaturi_100g": _clean_num_value(nt100.get("polinsaturi")),
            "polinsaturi_totale": _clean_num_value(nttot.get("polinsaturi")),
            "colesterolo_totale_100g": _clean_num_value(nt100.get("colesterolo")),
            "colesterolo_totale_totale": _clean_num_value(nttot.get("colesterolo")),
        }
    )

    nutrient_keys = [
        "energia_100g",
        "energia_totale",
        "energia_porzione",
        "energia_ricetta",
        "carboidrati_totali_100g",
        "carboidrati_totali_totale",
        "carboidrati_100g",
        "carboidrati_ricetta",
        "carboidrati_porzione",
        "di_cui_zuccheri_100g",
        "di_cui_zuccheri_totale",
        "zuccheri_100g",
        "zuccheri_ricetta",
        "zuccheri_porzione",
        "grassi_totali_100g",
        "grassi_totali_totale",
        "grassi_totali_porzione",
        "grassi_totali_ricetta",
        "di_cui_saturi_100g",
        "di_cui_saturi_totale",
        "grassi_saturi_100g",
        "grassi_saturi_ricetta",
        "grassi_saturi_porzione",
        "proteine_totali_100g",
        "proteine_totali_totale",
        "proteine_100g",
        "proteine_totale",
        "proteine_ricetta",
        "proteine_porzione",
        "fibre_100g",
        "fibre_totale",
        "fibre_ricetta",
        "fibre_porzione",
        "sodio_100g",
        "sodio_totale",
        "sodio_ricetta",
        "sodio_porzione",
        "sodio_100g_mg",
        "sodio_totale_mg",
        "sodio_ricetta_mg",
        "sodio_porzione_mg",
        "colesterolo_100g",
        "colesterolo_totale",
        "colesterolo_ricetta",
        "colesterolo_porzione",
        "colesterolo_100g_mg",
        "colesterolo_ricetta_mg",
        "colesterolo_porzione_mg",
        "grassi_monoinsaturi_100g",
        "grassi_monoinsaturi_totale",
        "grassi_monoinsaturi_ricetta",
        "grassi_monoinsaturi_porzione",
        "grassi_polinsaturi_100g",
        "grassi_polinsaturi_totale",
        "grassi_polinsaturi_ricetta",
        "grassi_polinsaturi_porzione",
        "monoinsaturi_100g",
        "monoinsaturi_totale",
        "polinsaturi_100g",
        "polinsaturi_totale",
        "colesterolo_totale_100g",
        "colesterolo_totale_totale",
    ]
    for key in nutrient_keys:
        val = _to_float(ctx.get(key))
        ctx[key] = val if val is not None else 0

    ctx["kcal_ricetta"] = _to_float(recipe.get("kcal_ricetta") or ctx.get("energia_ricetta"))
    ctx["kcal_per_porzione"] = _to_float(recipe.get("kcal_per_porzione") or recipe.get("kcal_porzione") or ctx.get("energia_porzione"))

    energia_100g_val = _clean_num_value(nt100.get("energia"))
    energia_tot_val = _clean_num_value(nttot.get("energia"))
    energia_por_val = _clean_num_value(ntpor.get("energia"))
    if energia_100g_val not in ("", None):
        ctx["energia_100g_kj"] = round(_num(energia_100g_val) * 4.184, 2)
    if energia_tot_val not in ("", None):
        ctx["energia_ricetta_kj"] = round(_num(energia_tot_val) * 4.184, 2)
    if energia_por_val not in ("", None):
        ctx["energia_porzione_kj"] = round(_num(energia_por_val) * 4.184, 2)

    cost_lines = recipe.get("cost_lines")
    if not isinstance(cost_lines, list) or not cost_lines:
        cost_lines = []
    cleaned_cost_lines: List[Dict[str, Any]] = []
    for row in cost_lines:
        if not isinstance(row, dict):
            continue
        clean_row: Dict[str, Any] = {}
        for k, v in row.items():
            clean_row[str(k)] = _clean_text_value(v)
        clean_row["nome"] = _clean_text_value(
            row.get("nome")
            or row.get("ingrediente")
            or row.get("ingredient")
            or row.get("name")
            or row.get("parsed_name")
            or row.get("line")
        )
        clean_row["name"] = clean_row.get("nome") or ""
        clean_row["parsed_name"] = _clean_text_value(row.get("parsed_name"))
        clean_row["line"] = _clean_text_value(row.get("line"))
        clean_row["quantita"] = _clean_text_value(
            row.get("quantita") or row.get("quantita_usata") or row.get("grams")
        )
        clean_row["quantita_raw"] = _clean_text_value(row.get("quantita_raw") or row.get("quantita_usata"))
        clean_row["grams"] = row.get("grams") or row.get("quantita")
        clean_row["unita"] = _clean_text_value(row.get("unita") or row.get("unit"))
        clean_row["unit"] = clean_row.get("unita") or ""
        clean_row["acquisto_minimo_g"] = _to_float(
            row.get("acquisto_minimo_g") or row.get("acquisto_minimo") or row.get("peso_min_acquisto")
        )
        clean_row["costo_ingrediente"] = _to_float(
            row.get("costo_ingrediente") or row.get("costo_usato") or row.get("prezzo_calcolato")
        )
        clean_row["costo_per_qty"] = _to_float(row.get("costo_per_qty"))
        clean_row["costo_usato"] = _to_float(row.get("costo_usato"))
        clean_row["costo_spesa"] = _to_float(row.get("costo_spesa") or row.get("prezzo_alimento_acquisto"))
        clean_row["costo_acquisto_minimo"] = _to_float(row.get("costo_acquisto_minimo"))
        if not clean_row.get("scarto"):
            for alt in ("scarto_pct", "waste_pct"):
                if clean_row.get(alt):
                    clean_row["scarto"] = clean_row.get(alt)
                    break
        clean_row["peso_min_acquisto"] = _normalize_unit_text(clean_row.get("peso_min_acquisto"))
        clean_row["quantita_usata"] = _normalize_unit_text(clean_row.get("quantita_usata"))
        clean_row["prezzo_kg_ud"] = _format_price_unit(clean_row.get("prezzo_kg_ud"))
        clean_row["prezzo_alimento_acquisto"] = _format_currency(clean_row.get("prezzo_alimento_acquisto"))
        clean_row["prezzo_calcolato"] = _format_currency(clean_row.get("prezzo_calcolato"))
        cleaned_cost_lines.append(clean_row)
    ctx["ingredienti_dettaglio"] = cleaned_cost_lines
    ctx["cost_rows"] = cleaned_cost_lines

    costo_totale_ricetta_val = _to_float(recipe.get("costo_totale_ricetta") or recipe.get("spesa_totale_ricetta"))
    costo_per_porzione_val = _to_float(recipe.get("costo_per_porzione") or recipe.get("spesa_per_porzione"))
    costo_spesa_totale = _to_float(recipe.get("costo_spesa_totale") or recipe.get("spesa_totale_acquisto"))
    costo_spesa_per_porzione = _to_float(recipe.get("costo_spesa_per_porzione"))
    costo_materia_usata = _to_float(recipe.get("costo_materia_usata"))
    prezzo_vendita = _to_float(recipe.get("prezzo_vendita"))
    food_cost_percent = _to_float(recipe.get("food_cost_percent") or recipe.get("food_cost"))
    ricarico_percent = _to_float(recipe.get("ricarico_percent"))
    margine_percent = _to_float(recipe.get("margine_percent"))
    prezzo_consigliato = _to_float(recipe.get("prezzo_consigliato"))
    fonte_prezzi = _clean_text_value(recipe.get("fonte_prezzi"))

    ctx["spesa_totale_ricetta"] = costo_totale_ricetta_val or 0
    ctx["costo_totale_ricetta"] = costo_totale_ricetta_val or 0
    ctx["spesa_per_porzione"] = costo_per_porzione_val or 0
    ctx["costo_per_porzione"] = costo_per_porzione_val or 0
    ctx["spesa_totale_acquisto"] = costo_spesa_totale or 0
    ctx["costo_spesa_totale"] = costo_spesa_totale or 0
    ctx["costo_spesa_per_porzione"] = costo_spesa_per_porzione or 0
    ctx["costo_materia_usata"] = costo_materia_usata or costo_totale_ricetta_val or 0
    ctx["prezzo_vendita"] = prezzo_vendita or 0
    ctx["food_cost_percent"] = food_cost_percent or 0
    ctx["food_cost"] = recipe.get("food_cost") or (f"{food_cost_percent:.1f}%" if food_cost_percent else "")
    ctx["ricarico_percent"] = ricarico_percent or 0
    ctx["margine_percent"] = margine_percent or 0
    ctx["prezzo_consigliato"] = prezzo_consigliato or 0
    ctx["fonte_prezzi"] = fonte_prezzi

    return ctx


# ------------------------------------------------------------
# Pipeline
# ------------------------------------------------------------
def process_single_file(filepath: str, options: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    res = analyze_files([filepath], options=options)
    if not res.get("ok"):
        return {"ok": False, "error": res.get("error", "Errore") or "Errore"}
    return {"ok": True, "recipe": res.get("recipe")}


def analyze_files(
    file_paths: List[str],
    progress: Any = None,
    options: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    opts = options or {}
    run_id = str(opts.get("run_id") or uuid.uuid4())
    file_id = str(opts.get("file_id") or (Path(file_paths[0]).stem if file_paths else "unknown"))
    logger = get_logger(run_id)

    subscription_tier = str(opts.get("subscription_tier") or "").lower().strip()
    allow_shared_cloud = bool(opts.get("allow_shared_cloud_ai", False))

    use_ai = bool(opts.get("use_ai", True))
    ai_complete_missing = bool(opts.get("ai_complete_missing", True))
    force_ai_full = bool(opts.get("force_ai_full", True))
    cloud_available = _cloud_ai_available(allow_shared=allow_shared_cloud, subscription_tier=subscription_tier)
    cloud_only = bool(cloud_available)
    cloud_fast = bool(cloud_available)

    steps: List[Tuple[str, float]] = [("extract", 30.0), ("parse", 10.0), ("ai_complete", 40.0), ("enrich", 20.0)]
    plan = _ProgressPlan(progress, steps)
    ai_completion_ok = False
    ai_error: Optional[str] = None
    ai_provider: Optional[str] = None
    t0 = time.monotonic()
    log_event(logger, run_id=run_id, file_id=file_id, stage="start", status="start", message=f"Start analyze {file_paths}")

    if not file_paths:
        msg = "Nessun file selezionato"
        log_event(logger, run_id=run_id, file_id=file_id, stage="start", status="error", message=msg)
        return {"ok": False, "error": msg}

    plan.set("extract", 5, "Estrazione testo...")

    def _ocr_progress(pct: int, msg: str) -> None:
        plan.set("extract", pct, msg)

    t_extract_start = time.monotonic()
    text, debug = extract_text_from_paths(file_paths, options=opts, progress_cb=_ocr_progress)
    log_event(
        logger,
        run_id=run_id,
        file_id=file_id,
        stage="extract",
        status="done",
        message=f"OCR done len={len(text)}",
        extra={"elapsed_s": round(time.monotonic() - t_extract_start, 3)},
    )
    if not text.strip():
        msg = "Testo vuoto"
        log_event(logger, run_id=run_id, file_id=file_id, stage="extract", status="error", message=msg)
        return {"ok": False, "error": msg}

    plan.done("extract", "Estrazione completata")
    plan.set("parse", 10, "Parsing testo...")
    recipe_raw, missing_raw = parse_recipe_text(text)

    recipe = recipe_raw
    missing = missing_raw

    plan.done("parse", "Parsing completato")
    log_event(
        logger,
        run_id=run_id,
        file_id=file_id,
        stage="parse",
        status="done",
        message="Parsing completato",
        extra={"elapsed_s": round(time.monotonic() - t0, 3)},
    )

    if not recipe.get("title") or recipe.get("title") == "Ricetta":
        recipe["title"] = _title_from_paths(file_paths)

    recipe["source_files"] = list(file_paths)

    if use_ai and ai_complete_missing:
        missing_fields = _collect_missing_fields(recipe, missing)
        if cloud_only or force_ai_full:
            missing_fields = sorted(set(missing_fields + _CLOUD_REQUIRED_KEYS))
        if not missing_fields and not force_ai_full:
            ai_completion_ok = True
        if missing_fields:
            plan.set("ai_complete", 10, "AI: integrazione dati...")
            ai_completion_error: Optional[str] = None
            ai_recipe, ai_provider, ai_error = _complete_missing_with_ai(
                recipe,
                text,
                missing_fields,
                allow_cloud=True,
                allow_ollama=False,
                retries_cloud=2,
                subscription_tier=subscription_tier,
                allow_shared_cloud=allow_shared_cloud,
            )
            if isinstance(ai_recipe, dict) and ai_recipe:
                applied = _apply_ai_patch(recipe, ai_recipe, allow_override_lists=True)
                if applied:
                    recipe["ai_completion"] = {"provider": ai_provider, "fields": applied}
                    ai_completion_ok = True
                missing = _collect_missing_fields(recipe, missing)
            if ai_error and not ai_completion_ok:
                ai_completion_error = ai_error
                ai_error = ai_error

            if not ai_completion_ok and not cloud_only and not cloud_available:
                missing_fields = _collect_missing_fields(recipe, missing)
                if force_ai_full:
                    missing_fields = sorted(set(missing_fields + _CLOUD_REQUIRED_KEYS))
                if missing_fields:
                    plan.set("ai_complete", 60, "AI: fallback locale...")
                    ai_recipe, ai_provider, ai_error = _complete_missing_with_ai(
                        recipe,
                        text,
                        missing_fields,
                        allow_cloud=False,
                        allow_ollama=False,
                        retries_ollama=2,
                        subscription_tier=subscription_tier,
                        allow_shared_cloud=allow_shared_cloud,
                    )
                    if isinstance(ai_recipe, dict) and ai_recipe:
                        applied = _apply_ai_patch(recipe, ai_recipe, allow_override_lists=True)
                        if applied:
                            recipe["ai_completion"] = {"provider": ai_provider, "fields": applied}
                            ai_completion_ok = True
                        missing = _collect_missing_fields(recipe, missing)
                    if ai_error and not ai_completion_ok:
                        ai_completion_error = ai_error

            if ai_completion_error:
                recipe["ai_completion_error"] = ai_completion_error
            else:
                recipe.pop("ai_completion_error", None)
            plan.done("ai_complete", "AI: integrazione completata")

    _apply_saverio_rules(recipe)

    if ai_cloud is not None and not cloud_available:
        try:
            if ai_cloud.should_call_cloud(recipe, missing):
                patch, meta = ai_cloud.request_patch(recipe, text, missing_fields=missing)
                if patch:
                    ai_cloud.apply_patch(recipe, patch)
                    recipe["cloud_ai"] = meta
        except Exception:
            pass

            log_event(
                logger,
                run_id=run_id,
                file_id=file_id,
                stage="ai_complete",
                status="done" if ai_completion_ok else "error",
                message="AI completata",
                extra={"provider": ai_provider or "cloud", "elapsed_s": round(time.monotonic() - t0, 3), "ai_error": ai_error},
            )

    if cloud_fast:
        missing_after_cloud = _collect_missing_fields(recipe, missing)
        needs_enrich = any(
            field in missing_after_cloud
            for field in ("costs", "nutrition", "allergens", "equipment", "diets")
        )
        if needs_enrich:
            allow_local_ai = False
            if allow_local_ai:
                plan.set("enrich", 10, "Arricchimento dati (fallback locale)")
            else:
                plan.set("enrich", 10, "Arricchimento dati (locale)")
            _enrich_data(recipe, source_text=text, allow_local_ai=allow_local_ai)
            plan.done("enrich", "Arricchimento dati completato")
        else:
            plan.set("enrich", 100, "Arricchimento dati (saltato: cloud)")
    else:
        plan.set("enrich", 10, "Arricchimento dati...")
        _enrich_data(recipe, source_text=text, allow_local_ai=False)
        plan.done("enrich", "Arricchimento dati completato")

    if use_ai and ai_complete_missing and not cloud_available:
        missing_fields = _collect_missing_fields(recipe, missing)
        if force_ai_full:
            missing_fields = sorted(set(missing_fields + _CLOUD_REQUIRED_KEYS))
        if missing_fields:
            plan.set("ai_complete", 10, "AI: completamento dati...")
            ai_recipe, ai_provider, ai_error = _complete_missing_with_ai(
                recipe,
                text,
                missing_fields,
                allow_cloud=True,
                allow_ollama=False,
                retries_ollama=1,
                subscription_tier=subscription_tier,
                allow_shared_cloud=allow_shared_cloud,
            )
            if isinstance(ai_recipe, dict) and ai_recipe:
                applied = _apply_ai_patch(recipe, ai_recipe, allow_override_lists=False)
                if applied:
                    recipe["ai_completion"] = {"provider": ai_provider, "fields": applied}
                _apply_saverio_rules(recipe)
                _enrich_data(recipe, source_text=text, allow_local_ai=False)
                missing = _collect_missing_fields(recipe, missing)
            if ai_error:
                recipe["ai_completion_error"] = ai_error
                log_event(
                    logger,
                    run_id=run_id,
                    file_id=file_id,
                    stage="ai_complete",
                    status="error",
                    message=f"AI fallback error: {ai_error}",
                )
            plan.done("ai_complete", "AI: completamento dati terminato")

    if use_ai and ai_complete_missing and cloud_available:
        template_ctx = build_template_context(recipe)
        missing_template = _collect_missing_template_fields(template_ctx)
        if missing_template:
            plan.set("ai_complete", 85, "AI: completamento campi template...")
            ai_recipe, provider, ai_error = _complete_missing_with_ai(
                recipe,
                text,
                list(_CLOUD_REQUIRED_KEYS),
                allow_cloud=True,
                allow_ollama=False,
                subscription_tier=subscription_tier,
                allow_shared_cloud=allow_shared_cloud,
            )
            if isinstance(ai_recipe, dict) and ai_recipe:
                applied = _apply_ai_patch(recipe, ai_recipe, allow_override_lists=True)
                if applied:
                    recipe["ai_completion"] = {"provider": provider, "fields": applied}
            if ai_error:
                recipe["ai_completion_error"] = ai_error
            _apply_saverio_rules(recipe)
            _enrich_data(recipe, source_text=text, allow_local_ai=False)
            template_ctx = build_template_context(recipe)
            missing_template = _collect_missing_template_fields(template_ctx)
            if missing_template:
                recipe["missing_fields"] = missing_template

    _p_set(progress, 100, "done", "Fatto")

    missing_filtered = _collect_missing_fields(recipe, missing)
    template_ctx = build_template_context(recipe)
    missing_template = _collect_missing_template_fields(template_ctx)
    if missing_template:
        missing_filtered = sorted(set(missing_filtered + missing_template))
        recipe["missing_fields"] = missing_filtered
    if ai_error:
        recipe["ai_completion_error"] = ai_error
    debug.update(
        {
            "ingredients_count": len(recipe.get("ingredients", []) if isinstance(recipe.get("ingredients"), list) else []),
            "steps_count": len(recipe.get("steps", []) if isinstance(recipe.get("steps"), list) else []),
            "servings": recipe.get("servings"),
            "title": recipe.get("title"),
            "ai_error": recipe.get("ai_completion_error"),
            "missing_filtered": missing_filtered,
        }
    )

    log_event(
        logger,
        run_id=run_id,
        file_id=file_id,
        stage="done",
        status="ok",
        message="Analyze complete",
        extra={"missing": len(missing_filtered), "ai_error": ai_error, "elapsed_s": round(time.monotonic() - t0, 3)},
    )

    return {
        "ok": True,
        "recipe": recipe,
        "missing_fields": missing_filtered,
        "debug": debug,
    }


# ------------------------------------------------------------
# Export helpers
# ------------------------------------------------------------
def list_templates() -> List[str]:
    if _get_pdf_templates:
        return _get_pdf_templates()
    return ["Template_Ricetta_AI"]


def render_html_template(recipe: Dict[str, Any], template: str, assets_dir: Optional[str] = None, silent_mode: bool = False) -> str:
    if not isinstance(recipe, dict):
        return ""

    use_as_context = bool(recipe.get("_preview_demo"))
    if use_as_context:
        ctx = dict(recipe)
    else:
        ctx = build_template_context(recipe)

    if _render_html_template:
        return _render_html_template(ctx, template=template, assets_dir=assets_dir, silent_mode=silent_mode)
    return ""


def _default_assets_dir() -> Optional[str]:
    base = project_root() / "templates" / "assets"
    if base.exists() and base.is_dir():
        return str(base)
    return None


def export_pdf(
    recipe: Dict[str, Any],
    out_path: str,
    template: str = "Template_Ricetta_AI",
    page_size: str = DEFAULT_EXPORT_PDF_SIZE,
    progress: Any = None,
) -> Dict[str, Any]:
    if not export_recipe_pdf:
        return {"ok": False, "error": "Modulo PDF non disponibile"}

    _p_set(progress, 2, "export", "Esportazione PDF...")
    ctx = build_template_context(recipe)
    assets_dir = _default_assets_dir()
    export_recipe_pdf(ctx, out_path, template=template, page_size=page_size, assets_dir=assets_dir)
    _p_set(progress, 100, "export", "PDF creato")
    return {"ok": True, "out_path": out_path}


def export_docx(
    recipe: Dict[str, Any],
    out_path: str,
) -> Dict[str, Any]:
    if not export_recipe_docx:
        return {"ok": False, "error": "Modulo DOCX non disponibile"}
    ctx = build_template_context(recipe)
    ok = export_recipe_docx(ctx, out_path)
    return {"ok": ok, "out_path": out_path}


# ------------------------------------------------------------
# Batch
# ------------------------------------------------------------
def _clean_title_for_filename(title: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", " ", title or "").strip()
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned or "Ricetta"


def run_pipeline(
    file_paths: List[str],
    progress_cb: Any = None,
    options: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    total = len(file_paths)
    last_valid_result: Dict[str, Any] = {}

    opts = options or {}
    template = str(opts.get("template") or opts.get("pdf_template") or "Template_Ricetta_AI")
    page_size = str(opts.get("page_size") or opts.get("pdf_page_size") or DEFAULT_EXPORT_PDF_SIZE)
    export_pdf_each = bool(opts.get("export_pdf", True))
    export_docx_each = bool(opts.get("export_docx", True))

    db = ArchiveDB.default() if ArchiveDB else None

    base_out = Path(__file__).resolve().parent.parent / "Output_Elaborati"
    base_out.mkdir(parents=True, exist_ok=True)

    for i, fpath in enumerate(file_paths):
        fname = Path(fpath).name
        _p_set(progress_cb, int((i / max(1, total)) * 100), "work", f"Elaboro: {fname}")

        res = process_single_file(fpath, options=opts)
        if not res.get("ok"):
            print(f"Skipped {fname}: {res.get('error')}")
            continue

        recipe = res.get("recipe") or {}
        last_valid_result = res

        if db:
            try:
                db.save_recipe(recipe)
            except Exception as e:
                print(f"DB Error: {e}")

        clean_title = _clean_title_for_filename(recipe.get("title", "Ricetta"))
        category = recipe.get("category", "Altro")

        cat_dir = base_out / str(category)
        cat_dir.mkdir(parents=True, exist_ok=True)

        base_filename = cat_dir / clean_title

        if export_pdf_each and export_recipe_pdf:
            try:
                pdf_name = f"{base_filename}.pdf"
                export_recipe_pdf(
                    build_template_context(recipe),
                    pdf_name,
                    template=template,
                    page_size=page_size,
                    assets_dir=_default_assets_dir(),
                )
            except Exception as e:
                print(f"PDF Export Error: {e}")

        if export_docx_each and export_recipe_docx:
            try:
                docx_name = f"{base_filename}.docx"
                export_recipe_docx(build_template_context(recipe), docx_name)
            except Exception as e:
                print(f"Word Export Error: {e}")

    _p_set(progress_cb, 100, "done", "Finito! Controlla Output_Elaborati")
    return last_valid_result
